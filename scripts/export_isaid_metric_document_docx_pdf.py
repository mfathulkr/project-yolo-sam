from __future__ import annotations

import argparse
import sys
from pathlib import Path

import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from matplotlib.backends.backend_pdf import PdfPages
from PIL import Image

ROOT = Path(__file__).resolve().parents[1]
SCRIPTS = ROOT / "scripts"
if str(SCRIPTS) not in sys.path:
    sys.path.insert(0, str(SCRIPTS))

from write_isaid_metric_document import (  # noqa: E402
    CONTEXT_BULLETS,
    CONTEXT_HEADING,
    DETECTOR_METRIC_NOTE,
    DETECTOR_HEADING,
    DISCUSSION_HEADING,
    DOCUMENT_TITLE,
    FINDINGS_HEADING,
    METRIC_BULLETS,
    METRIC_LOGIC_HEADING,
    QUALITATIVE_HEADING,
    QUALITATIVE_NOTE,
    SCOPE_HEADING,
    SCOPE_BULLETS,
    SEGMENTATION_HEADING,
    SEGMENTATION_TABLES_NOTE,
    TABLES,
    discussion_lines,
    display_detector_metrics,
    display_summary,
    table_proxy_note,
)


SUCCESS_COLUMNS = {
    "Avg IoU",
    "Avg Dice",
    "Avg Precision",
    "Avg Recall",
    "mAP50 proxy",
    "mAP75 proxy",
    "mAP90 proxy",
    "mAP50-95 proxy",
    "BBox mAP50",
    "BBox mAP75",
    "BBox mAP90",
    "BBox mAP50-95",
    "BBox Precision@0.50",
    "BBox Recall@0.50",
    "BBox Precision@0.75",
    "BBox Recall@0.75",
    "BBox Precision@0.90",
    "BBox Recall@0.90",
}

QUALITATIVE_IMAGES = [
    (
        "No Overlap / Low Mask Area",
        ROOT / "presentation_isaid_vehicle_sam3_sam2_study" / "figures" / "sample_cases" / "no_overlap__low_mask_area__P2766_0016_hero.png",
    ),
    (
        "No Overlap / High Mask Area",
        ROOT / "presentation_isaid_vehicle_sam3_sam2_study" / "figures" / "sample_cases" / "no_overlap__high_mask_area__P0199_0002_hero.png",
    ),
    (
        "Overlap / Low Mask Area",
        ROOT / "presentation_isaid_vehicle_sam3_sam2_study" / "figures" / "sample_cases" / "overlap__low_mask_area__P2404_0002_hero.png",
    ),
    (
        "Overlap / High Mask Area",
        ROOT / "presentation_isaid_vehicle_sam3_sam2_study" / "figures" / "sample_cases" / "overlap__high_mask_area__P2781_0005_hero.png",
    ),
]


def parse_args() -> argparse.Namespace:
    base_dir = ROOT / "presentation_isaid_vehicle_sam3_sam2_study"
    parser = argparse.ArgumentParser(description="iSAID metrik raporunu renkli DOCX ve PDF olarak dışa aktar.")
    parser.add_argument("--output-dir", type=Path, default=base_dir)
    parser.add_argument("--tables-dir", type=Path, default=base_dir / "tables" / "full_metric_document")
    parser.add_argument(
        "--detector-metrics",
        type=Path,
        default=ROOT / "results" / "isaid_vehicle_detector_metrics" / "yolo_detector_eval_metrics.csv",
    )
    return parser.parse_args()


def interpolate_color(value: float) -> str:
    value = max(0.0, min(1.0, value))
    red = (248, 105, 107)
    yellow = (255, 235, 132)
    green = (99, 190, 123)
    if value <= 0.5:
        t = value / 0.5
        left, right = red, yellow
    else:
        t = (value - 0.5) / 0.5
        left, right = yellow, green
    rgb = tuple(round(left[index] + (right[index] - left[index]) * t) for index in range(3))
    return "".join(f"{channel:02X}" for channel in rgb)


def numeric_value(text: object) -> float | None:
    try:
        value = float(str(text))
    except ValueError:
        return None
    if 0.0 <= value <= 1.0:
        return value
    return None


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: object, font_size: float = 7.0, bold: bool = False, color: str | None = None) -> None:
    cell.text = str(text)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.size = Pt(font_size)
            run.font.bold = bold
            if color:
                run.font.color.rgb = RGBColor.from_string(color)


def add_bullets(doc: Document, bullets: list[str]) -> None:
    for bullet in bullets:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.add_run(bullet).font.size = Pt(10)


def add_docx_table(doc: Document, df: pd.DataFrame, title: str | None = None, note: str | None = None) -> None:
    if title:
        doc.add_heading(title, level=2)
    if note:
        paragraph = doc.add_paragraph()
        paragraph.add_run(note).font.size = Pt(9)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for index, column in enumerate(df.columns):
        cell = table.rows[0].cells[index]
        set_cell_text(cell, column, font_size=7.0, bold=True, color="FFFFFF")
        shade_cell(cell, "1F4E79")

    for _, row in df.iterrows():
        cells = table.add_row().cells
        for index, column in enumerate(df.columns):
            value = row[column]
            cell = cells[index]
            set_cell_text(cell, value, font_size=6.2)
            if column in SUCCESS_COLUMNS:
                numeric = numeric_value(value)
                if numeric is not None:
                    shade_cell(cell, interpolate_color(numeric))


def load_summaries(tables_dir: Path) -> dict[str, pd.DataFrame]:
    raw_tables_dir = tables_dir / "raw_summaries"
    if not raw_tables_dir.exists():
        raw_tables_dir = tables_dir
    summaries: dict[str, pd.DataFrame] = {}
    for table_key, _, _ in TABLES:
        summaries[table_key] = pd.read_csv(raw_tables_dir / f"summary_{table_key}.csv")
    return summaries


def build_docx(output_path: Path, summaries: dict[str, pd.DataFrame], detector_table: pd.DataFrame) -> None:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    for attr in ["top_margin", "bottom_margin", "left_margin", "right_margin"]:
        setattr(section, attr, Inches(0.35))

    title = doc.add_heading(DOCUMENT_TITLE, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading(SCOPE_HEADING, level=1)
    add_bullets(doc, SCOPE_BULLETS)

    doc.add_heading(METRIC_LOGIC_HEADING, level=1)
    add_bullets(doc, METRIC_BULLETS)

    doc.add_heading(CONTEXT_HEADING, level=1)
    add_bullets(doc, CONTEXT_BULLETS)

    doc.add_heading(DETECTOR_HEADING, level=1)
    paragraph = doc.add_paragraph()
    paragraph.add_run(DETECTOR_METRIC_NOTE).font.size = Pt(10)
    add_docx_table(doc, detector_table)

    doc.add_heading(SEGMENTATION_HEADING, level=1)
    doc.add_paragraph(SEGMENTATION_TABLES_NOTE)
    for table_key, title_text, _ in TABLES:
        add_docx_table(
            doc,
            display_summary(summaries[table_key]),
            title=title_text,
            note=table_proxy_note(summaries[table_key]),
        )

    doc.add_heading(QUALITATIVE_HEADING, level=1)
    doc.add_paragraph(QUALITATIVE_NOTE)
    for title_text, image_path in QUALITATIVE_IMAGES:
        doc.add_heading(title_text, level=2)
        doc.add_picture(str(image_path), width=Inches(9.8))

    doc.add_heading(DISCUSSION_HEADING, level=1)
    add_bullets(doc, discussion_lines(summaries))
    doc.save(output_path)


def add_pdf_text_page(pdf: PdfPages, title: str, sections: list[tuple[str, list[str]]]) -> None:
    fig = plt.figure(figsize=(11.69, 8.27))
    fig.patch.set_facecolor("white")
    plt.axis("off")
    y = 0.95
    fig.text(0.05, y, title, fontsize=20, weight="bold", va="top")
    y -= 0.08
    for section_title, bullets in sections:
        fig.text(0.05, y, section_title, fontsize=13, weight="bold", va="top")
        y -= 0.035
        for bullet in bullets:
            wrapped = wrap_text("- " + bullet, 145)
            fig.text(0.07, y, wrapped, fontsize=8.2, va="top")
            y -= 0.028 * (wrapped.count("\n") + 1)
            if y < 0.05:
                pdf.savefig(fig, bbox_inches="tight")
                plt.close(fig)
                fig = plt.figure(figsize=(11.69, 8.27))
                fig.patch.set_facecolor("white")
                plt.axis("off")
                y = 0.95
        y -= 0.02
    pdf.savefig(fig, bbox_inches="tight")
    plt.close(fig)


def wrap_text(text: str, width: int) -> str:
    words = text.split()
    lines: list[str] = []
    current: list[str] = []
    current_len = 0
    for word in words:
        extra = 1 if current else 0
        if current_len + len(word) + extra > width:
            lines.append(" ".join(current))
            current = [word]
            current_len = len(word)
        else:
            current.append(word)
            current_len += len(word) + extra
    if current:
        lines.append(" ".join(current))
    return "\n".join(lines)


def add_pdf_table(pdf: PdfPages, df: pd.DataFrame, title: str, note: str | None = None) -> None:
    fig, ax = plt.subplots(figsize=(17.0, 9.5))
    fig.patch.set_facecolor("white")
    ax.axis("off")
    ax.set_title(title, fontsize=16, weight="bold", pad=34 if note else 16)
    if note:
        fig.text(0.5, 0.915, wrap_text(note, 185), ha="center", va="top", fontsize=8.0)

    table = ax.table(
        cellText=df.values,
        colLabels=df.columns,
        cellLoc="center",
        loc="center",
        colColours=["#1F4E79"] * len(df.columns),
    )
    table.auto_set_font_size(False)
    table.set_fontsize(6.2)
    table.scale(1.0, 1.55)

    for (row_index, col_index), cell in table.get_celld().items():
        cell.set_edgecolor("#666666")
        cell.set_linewidth(0.35)
        if row_index == 0:
            cell.set_text_props(color="white", weight="bold")
            continue
        column = df.columns[col_index]
        if column in SUCCESS_COLUMNS:
            numeric = numeric_value(df.iloc[row_index - 1, col_index])
            if numeric is not None:
                cell.set_facecolor("#" + interpolate_color(numeric))

    pdf.savefig(fig, bbox_inches="tight")
    plt.close(fig)


def add_pdf_image_page(pdf: PdfPages, title: str, image_path: Path) -> None:
    with Image.open(image_path) as image:
        width, height = image.size
        fig_width = 11.69
        fig_height = max(8.27, fig_width * height / width + 0.8)
        fig, ax = plt.subplots(figsize=(fig_width, fig_height))
        fig.patch.set_facecolor("white")
        ax.axis("off")
        ax.set_title(title, fontsize=15, weight="bold", pad=12)
        ax.imshow(image)
        pdf.savefig(fig, bbox_inches="tight")
        plt.close(fig)


def build_pdf(output_path: Path, summaries: dict[str, pd.DataFrame], detector_table: pd.DataFrame) -> None:
    with PdfPages(output_path) as pdf:
        add_pdf_text_page(
            pdf,
            DOCUMENT_TITLE,
            [
                (SCOPE_HEADING, SCOPE_BULLETS),
                (METRIC_LOGIC_HEADING, METRIC_BULLETS),
                (CONTEXT_HEADING, CONTEXT_BULLETS),
            ],
        )
        add_pdf_table(pdf, detector_table, DETECTOR_HEADING, note=DETECTOR_METRIC_NOTE)
        for table_key, title_text, _ in TABLES:
            add_pdf_table(
                pdf,
                display_summary(summaries[table_key]),
                title_text,
                note=table_proxy_note(summaries[table_key]),
            )
        for title_text, image_path in QUALITATIVE_IMAGES:
            add_pdf_image_page(pdf, title_text, image_path)
        add_pdf_text_page(pdf, DISCUSSION_HEADING, [(FINDINGS_HEADING, discussion_lines(summaries))])


def main() -> None:
    args = parse_args()
    args.output_dir.mkdir(parents=True, exist_ok=True)
    summaries = load_summaries(args.tables_dir)
    detector_table = display_detector_metrics(args.detector_metrics)
    if detector_table.empty:
        raise FileNotFoundError(f"Detector metrics not found or empty: {args.detector_metrics}")

    docx_path = args.output_dir / "isaid_vehicle_full_metric_document_colored.docx"
    pdf_path = args.output_dir / "isaid_vehicle_full_metric_document_colored.pdf"
    build_docx(docx_path, summaries, detector_table)
    build_pdf(pdf_path, summaries, detector_table)
    print(docx_path)
    print(pdf_path)


if __name__ == "__main__":
    main()
