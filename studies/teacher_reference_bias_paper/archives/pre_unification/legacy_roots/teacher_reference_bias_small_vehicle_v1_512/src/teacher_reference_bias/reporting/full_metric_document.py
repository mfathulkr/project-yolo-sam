from __future__ import annotations

import hashlib
import json
import math
import os
import re
import textwrap
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

import matplotlib

matplotlib.use("Agg")

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from matplotlib.backends.backend_pdf import PdfPages
from PIL import Image, ImageDraw, ImageFont


REPO_ROOT = Path(__file__).resolve().parents[5]

STRATA = (
    ("overall", "Overall"),
    ("no_overlap__low_mask_area", "No Overlap × Low Mask Area"),
    ("no_overlap__high_mask_area", "No Overlap × High Mask Area"),
    ("overlap__low_mask_area", "Overlap × Low Mask Area"),
    ("overlap__high_mask_area", "Overlap × High Mask Area"),
)

METRICS = (
    ("mean_iou", "Avg IoU"),
    ("mean_dice", "Avg Dice"),
    ("mean_precision", "Avg Precision"),
    ("mean_recall", "Avg Recall"),
    ("success_at_iou_50", "IoU ≥ 0.50"),
    ("success_at_iou_75", "IoU ≥ 0.75"),
    ("success_at_iou_90", "IoU ≥ 0.90"),
)

SUCCESS_COLUMNS = {display_name for _, display_name in METRICS} | {
    "Human IoU",
    "SAM1 Pseudo IoU",
    "Reference IoU",
    "BBox mAP50",
    "BBox mAP75",
    "BBox mAP90",
    "BBox mAP50-95",
    "BBox Precision@0.50",
    "BBox Recall@0.50",
    "BBox Precision@0.75",
    "BBox Recall@0.75",
    "BBox Precision@0.90",
    "BBox Recall@0.90",
}

IMAGE_COUNTS = {
    "overall": 512,
    "no_overlap__low_mask_area": 128,
    "no_overlap__high_mask_area": 128,
    "overlap__low_mask_area": 128,
    "overlap__high_mask_area": 128,
}

QUALITATIVE_LABELS = (
    ("no_overlap__low_mask_area", "No Overlap / Low Mask Area"),
    ("no_overlap__high_mask_area", "No Overlap / High Mask Area"),
    ("overlap__low_mask_area", "Overlap / Low Mask Area"),
    ("overlap__high_mask_area", "Overlap / High Mask Area"),
)


@dataclass(frozen=True)
class ReferenceSection:
    reference_type: str
    title: str
    note: str


@dataclass(frozen=True)
class ReportSpec:
    study_id: str
    dataset_id: str
    slug: str
    title: str
    dataset_label: str
    reference_sections: tuple[ReferenceSection, ...]
    qualitative_image: Path
    scope_bullets: tuple[str, ...]
    context_bullets: tuple[str, ...]
    discussion_bullets: tuple[str, ...]
    target_label: str = "küçük araç"


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def portable_manifest_path(path: Path) -> str:
    try:
        return path.resolve().relative_to(REPO_ROOT).as_posix()
    except ValueError as exc:
        raise ValueError(f"Report dependency is outside the repository: {path}") from exc


def interpolate_color(value: float) -> str:
    value = max(0.0, min(1.0, value))
    red = (248, 105, 107)
    yellow = (255, 235, 132)
    green = (99, 190, 123)
    if value <= 0.5:
        ratio = value / 0.5
        left, right = red, yellow
    else:
        ratio = (value - 0.5) / 0.5
        left, right = yellow, green
    rgb = tuple(
        round(left[index] + (right[index] - left[index]) * ratio)
        for index in range(3)
    )
    return "".join(f"{channel:02X}" for channel in rgb)


def numeric_value(value: object) -> float | None:
    match = re.match(r"^\s*([0-9]+(?:\.[0-9]+)?)", str(value))
    if match is None:
        return None
    number = float(match.group(1))
    return number if 0.0 <= number <= 1.0 else None


def ranking_comparison_sentence(
    human_rank: str,
    pseudo_rank: str,
) -> str:
    """Describe reference sensitivity without claiming a nonexistent rank change."""
    if human_rank == pseudo_rank:
        return (
            "Tam GT-bbox sıralaması her iki referansta da "
            f"{human_rank} biçiminde korunmuştur; buna rağmen skor "
            "düzeylerindeki değişim pseudo referansın ölçülen başarı "
            "büyüklüğünü etkilediğini gösterir."
        )
    return (
        "Tam GT-bbox sıralaması insan referansında "
        f"{human_rank}, SAM1 pseudo referansında {pseudo_rank} biçimindedir; "
        "sıralamadaki değişim pseudo referansın model seçimini "
        "etkileyebileceği riskini gösterir."
    )


def _format_metric(mean: float, std: float | None = None) -> str:
    if std is None or not math.isfinite(std):
        return f"{mean:.3f}"
    return f"{mean:.3f} ± {std:.3f}"


def _seed_note(seed_ids: tuple[int, ...]) -> str:
    if len(seed_ids) == 1:
        return f"sabit seed {seed_ids[0]} sonucudur"
    joined = ", ".join(str(seed) for seed in seed_ids)
    return f"seed {joined} ortalaması ± standart sapmasıdır"


def build_segmentation_table(
    aggregates: pd.DataFrame,
    *,
    dataset_id: str,
    reference_type: str,
    stratum: str,
) -> pd.DataFrame:
    selected = aggregates[
        (aggregates["dataset_id"] == dataset_id)
        & (aggregates["reference_type"] == reference_type)
        & (aggregates["stratum"] == stratum)
    ].copy()
    if selected.empty:
        raise ValueError(
            f"Missing aggregate rows for {dataset_id}/{reference_type}/{stratum}"
        )

    rows: list[dict[str, object]] = []
    yolo_seed_sets: list[tuple[int, ...]] = []
    for model in ("sam1", "sam2", "sam3"):
        for bbox_source in ("gt_bbox", "yolo_bbox"):
            pipeline = selected[
                (selected["model"] == model)
                & (selected["bbox_source"] == bbox_source)
            ].copy()
            if pipeline.empty:
                raise ValueError(
                    "Missing pipeline rows for "
                    f"{dataset_id}/{reference_type}/{stratum}/{model}/{bbox_source}"
                )

            seed_count = pipeline["detector_seed"].notna().sum()
            if bbox_source == "gt_bbox":
                if len(pipeline) != 1:
                    raise ValueError(
                        f"Expected one GT-bbox row, found {len(pipeline)}"
                    )
            else:
                seeds = tuple(
                    sorted(pipeline["detector_seed"].dropna().astype(int).unique())
                )
                if not seeds or len(pipeline) != seed_count or len(pipeline) != len(seeds):
                    raise ValueError(
                        "Expected one aggregate row per selected YOLO-bbox seed, "
                        f"found {list(seeds)} in {len(pipeline)} rows"
                    )
                yolo_seed_sets.append(seeds)

            counts = pipeline["instance_count"].astype(int).unique()
            scenes = pipeline["source_scene_count"].astype(int).unique()
            if len(counts) != 1 or len(scenes) != 1:
                raise ValueError(
                    f"Inconsistent counts for {dataset_id}/{model}/{bbox_source}"
                )

            row: dict[str, object] = {
                "Pipeline": (
                    f"{model.upper()} "
                    f"{'GT bbox' if bbox_source == 'gt_bbox' else 'YOLO bbox'}"
                ),
                "Images": IMAGE_COUNTS[stratum],
            }
            for metric_name, display_name in METRICS:
                mean = float(pipeline[metric_name].mean())
                std = (
                    float(pipeline[metric_name].std(ddof=1))
                    if bbox_source == "yolo_bbox" and len(pipeline) > 1
                    else None
                )
                row[display_name] = _format_metric(mean, std)
            rows.append(row)

    table = pd.DataFrame(rows)
    instance_counts = selected["instance_count"].astype(int).unique()
    if len(instance_counts) != 1:
        raise ValueError(
            f"Inconsistent instance counts for {dataset_id}/{reference_type}/{stratum}"
        )
    table.attrs["instance_count"] = int(instance_counts[0])
    if not yolo_seed_sets or len(set(yolo_seed_sets)) != 1:
        raise ValueError("YOLO-bbox pipelines do not use one consistent seed set")
    table.attrs["detector_seeds"] = yolo_seed_sets[0]
    return table


def segmentation_table_note(
    *,
    section: ReferenceSection,
    stratum: str,
    frame: pd.DataFrame,
    target_label: str = "küçük araç",
) -> str:
    image_count = IMAGE_COUNTS[stratum]
    instance_count = int(frame.attrs["instance_count"])
    formatted_instance_count = f"{instance_count:,}".replace(",", ".")
    return (
        f"Referans: {section.title}. Bu tablo {image_count} görüntüdeki "
        f"{formatted_instance_count} {target_label} örneğini kapsar. "
        f"YOLO bbox değerleri {_seed_note(tuple(frame.attrs['detector_seeds']))}."
    )


def build_detector_table(
    detector_summary: pd.DataFrame, *, dataset_id: str
) -> pd.DataFrame:
    selected = detector_summary[detector_summary["dataset_id"] == dataset_id]
    if len(selected) != 1:
        raise ValueError(
            f"Expected one detector summary for {dataset_id}, found {len(selected)}"
        )
    row = selected.iloc[0]
    seed_ids = tuple(
        int(value)
        for value in str(row.get("seed_ids", "")).split(",")
        if value.strip()
    )
    if not seed_ids:
        raise ValueError(f"Detector seed ids are missing for {dataset_id}")
    mapping = (
        ("bbox_AP50", "BBox mAP50"),
        ("bbox_AP75", "BBox mAP75"),
        ("bbox_AP90", "BBox mAP90"),
        ("bbox_AP50_95", "BBox mAP50-95"),
        ("precision_at_bbox_iou50", "BBox Precision@0.50"),
        ("recall_at_bbox_iou50", "BBox Recall@0.50"),
        ("precision_at_bbox_iou75", "BBox Precision@0.75"),
        ("recall_at_bbox_iou75", "BBox Recall@0.75"),
        ("precision_at_bbox_iou90", "BBox Precision@0.90"),
        ("recall_at_bbox_iou90", "BBox Recall@0.90"),
    )
    output: dict[str, object] = {
        "Detector": (
            f"YOLO26x (seed {seed_ids[0]})"
            if len(seed_ids) == 1
            else f"YOLO26x ({len(seed_ids)} seed)"
        ),
        "Images": 512,
    }
    for source_prefix, display_name in mapping:
        output[display_name] = _format_metric(
            float(row[f"{source_prefix}_mean"]),
            (
                float(row[f"{source_prefix}_std"])
                if len(seed_ids) > 1
                else None
            ),
        )
    table = pd.DataFrame([output])
    table.attrs["detector_seeds"] = seed_ids
    return table


def build_isaid_reference_effect_table(
    aggregates: pd.DataFrame,
) -> pd.DataFrame:
    selected = aggregates[
        (aggregates["dataset_id"] == "isaid_small_vehicle")
        & (aggregates["stratum"] == "overall")
    ]
    rows: list[dict[str, object]] = []
    for model in ("sam1", "sam2", "sam3"):
        for bbox_source in ("gt_bbox", "yolo_bbox"):
            condition = selected[
                (selected["model"] == model)
                & (selected["bbox_source"] == bbox_source)
            ]
            human = condition[condition["reference_type"] == "human"][
                "mean_iou"
            ].mean()
            pseudo = condition[condition["reference_type"] == "pseudo_sam1"][
                "mean_iou"
            ].mean()
            if pd.isna(human) or pd.isna(pseudo):
                raise ValueError(f"Missing iSAID dual-reference row for {model}")
            rows.append(
                {
                    "Model": model.upper(),
                    "BBox": "GT bbox" if bbox_source == "gt_bbox" else "YOLO bbox",
                    "Human IoU": f"{human:.3f}",
                    "SAM1 Pseudo IoU": f"{pseudo:.3f}",
                    "IoU Artışı": f"{pseudo - human:+.3f}",
                }
            )
    return pd.DataFrame(rows)


def build_samrs_shared_reference_table(
    shared_summary: pd.DataFrame,
    inflation_ci: list[dict[str, object]],
) -> pd.DataFrame:
    overall = shared_summary[shared_summary["stratum"] == "overall"]
    ci_by_model = {str(row["model"]): row for row in inflation_ci}
    rows: list[dict[str, object]] = []
    for model in ("sam1", "sam2", "sam3"):
        model_rows = overall[overall["model"] == model]
        human = model_rows[model_rows["reference_type"] == "human"]
        pseudo = model_rows[model_rows["reference_type"] == "pseudo_sam1"]
        if len(human) != 1 or len(pseudo) != 1:
            raise ValueError(f"Missing SAMRS shared-reference rows for {model}")
        ci = ci_by_model[model]
        rows.append(
            {
                "Model": model.upper(),
                "Human IoU": f"{float(human.iloc[0]['mean_iou']):.3f}",
                "SAM1 Pseudo IoU": f"{float(pseudo.iloc[0]['mean_iou']):.3f}",
                "IoU Artışı": f"{float(ci['estimate']):+.3f}",
                "%95 Güven Aralığı": (
                    f"[{float(ci['lower']):.3f}, "
                    f"{float(ci['upper']):.3f}]"
                ),
            }
        )
    return pd.DataFrame(rows)


def metric_bullets(
    target_label: str = "küçük araç",
    detector_seeds: tuple[int, ...] = (42,),
) -> tuple[str, ...]:
    return (
        "TP, modelin doğru biçimde nesne olarak işaretlediği pikseldir. FP, nesne olmadığı hâlde nesne diye işaretlenen; FN ise nesne olduğu hâlde kaçırılan pikseldir.",
        "IoU = TP / (TP + FP + FN). Tahmin ve referans maskenin ortak alanını birleşim alanına böler; 1 kusursuz, 0 hiç örtüşme yok demektir.",
        "Dice = 2TP / (2TP + FP + FN). IoU ile aynı davranışı farklı ölçekle ifade eder.",
        "Precision = TP / (TP + FP). Modelin boyadığı piksellerin ne kadarının gerçekten nesne olduğunu gösterir; fazla alan boyamak precision değerini düşürür.",
        "Recall = TP / (TP + FN). Gerçek nesne piksellerinin ne kadarının yakalandığını gösterir; eksik maske recall değerini düşürür.",
        f"Dört ortalama maske metriği nesne örneği düzeyinde (instance-level) önce her {target_label} için hesaplanır, sonra bütün örnekler eşit ağırlıkla ortalanır. Büyük nesneler küçük nesnelerin sonucunu perdelemez.",
        f"Her satırda insan anotasyonuyla varlığı bilinen bir {target_label} vardır. Bu nedenle boş pseudo referans eksik etikettir; tahmin de boş olsa bile maske metrikleri 0 kabul edilir ve referans kapsama kaybı ayrıca raporlanır.",
        f"IoU ≥ 0.50/0.75/0.90 sütunları, ilgili IoU eşiğini geçen {target_label} maskelerinin oranıdır. Bunlar mAP değildir ve raporda mAP gibi adlandırılmaz.",
        f"YOLO'nun kaçırdığı bir gerçek {target_label}, YOLO-bbox maske tablosunda boş tahmin olarak değerlendirilir ve o örneğin maske skorları sıfır olur. Herhangi bir gerçek nesneyle eşleşmeyen yanlış pozitif YOLO kutuları ise instance maske ortalamasına sahte bir referans örneği olarak eklenmez; bunların etkisi detector Precision, Recall ve mAP değerlerinde ölçülür.",
        f"Maske tabloları her GT {target_label} örneğini değerlendirir; YOLO'nun eşleştiremediği GT örnekleri de boş tahmin ve sıfır skorla hesaba katılır. Bu değerlendirme gerçek COCO segmentation AP ile aynı değildir. Confidence sırasındaki bütün maskeleri ve yanlış pozitifleri kullanan uçtan uca COCO mask AP bu raporda ayrıca çalıştırılmadığı için IoU eşik oranları AP veya mAP diye yeniden adlandırılmamıştır.",
        "Overall tablosu 512 görüntüyü, diğer tabloların her biri 128 görüntüyü kapsar.",
        f"GT-bbox satırları tek sabit koşuldur. YOLO-bbox satırlarındaki değerler {_seed_note(detector_seeds)}.",
    )


def detector_bullets(detector_seeds: tuple[int, ...] = (42,)) -> tuple[str, ...]:
    return (
        "Bu tablo yalnız YOLO detector kutularını değerlendirir; burada ölçülen bbox başarısıdır, maske başarısı değildir.",
        "BBox mAP50/mAP75/mAP90, tahmin kutusunun GT kutuyla sırasıyla en az 0,50/0,75/0,90 IoU yaptığı eşiklerde confidence sıralaması boyunca hesaplanan gerçek average precision değeridir.",
        "BBox mAP50-95, 0,50 ile 0,95 arasındaki on bbox IoU eşiğinin AP ortalamasıdır.",
        "BBox Precision ve Recall değerleri, doğrulama kümesinde seçilip testten önce sabitlenen güven eşiğinde hesaplanır.",
        f"Tablodaki değerler {_seed_note(detector_seeds)}.",
    )


def _shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_text(
    cell,
    value: object,
    *,
    font_size: float,
    bold: bool = False,
    color: str | None = None,
) -> None:
    cell.text = str(value)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.name = "DejaVu Sans"
            run.font.size = Pt(font_size)
            run.font.bold = bold
            if color:
                run.font.color.rgb = RGBColor.from_string(color)


def _add_docx_bullets(document: Document, bullets: Iterable[str]) -> None:
    for bullet in bullets:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(str(bullet))
        run.font.name = "DejaVu Sans"
        run.font.size = Pt(9)


def _add_docx_table(
    document: Document,
    frame: pd.DataFrame,
    *,
    title: str | None = None,
    note: str | None = None,
) -> None:
    if title:
        document.add_heading(title, level=2)
    if note:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(note)
        run.font.name = "DejaVu Sans"
        run.font.size = Pt(8.5)

    table = document.add_table(rows=1, cols=len(frame.columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for column_index, column in enumerate(frame.columns):
        cell = table.rows[0].cells[column_index]
        _set_cell_text(
            cell,
            column,
            font_size=6.1,
            bold=True,
            color="FFFFFF",
        )
        _shade_cell(cell, "1F4E79")

    for _, row in frame.iterrows():
        cells = table.add_row().cells
        for column_index, column in enumerate(frame.columns):
            value = row[column]
            cell = cells[column_index]
            _set_cell_text(cell, value, font_size=5.8)
            if column in SUCCESS_COLUMNS:
                numeric = numeric_value(value)
                if numeric is not None:
                    _shade_cell(cell, interpolate_color(numeric))


def _configure_docx(document: Document) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    for margin in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, margin, Inches(0.35))

    styles = document.styles
    styles["Normal"].font.name = "DejaVu Sans"
    styles["Normal"].font.size = Pt(9)
    for style_name in ("Title", "Heading 1", "Heading 2"):
        styles[style_name].font.name = "DejaVu Sans"


def _load_font(size: int, *, bold: bool = False) -> ImageFont.FreeTypeFont:
    name = "DejaVuSans-Bold.ttf" if bold else "DejaVuSans.ttf"
    path = Path("/usr/share/fonts/truetype/dejavu") / name
    return ImageFont.truetype(str(path), size=size)


def _contiguous_runs(indices: np.ndarray) -> list[tuple[int, int]]:
    if indices.size == 0:
        return []
    groups = np.split(indices, np.flatnonzero(np.diff(indices) > 1) + 1)
    return [(int(group[0]), int(group[-1]) + 1) for group in groups]


def _detect_qualitative_panel_ranges(
    source: Image.Image,
) -> tuple[list[tuple[int, int]], list[tuple[int, int]]]:
    pixels = np.asarray(source.convert("RGB"))
    height, width = pixels.shape[:2]
    nonwhite = np.any(pixels < 250, axis=2)
    row_ranges = [
        bounds
        for bounds in _contiguous_runs(
            np.flatnonzero(nonwhite.sum(axis=1) > 0.55 * width)
        )
        if bounds[1] - bounds[0] > 0.10 * height
    ]
    column_ranges = [
        bounds
        for bounds in _contiguous_runs(
            np.flatnonzero(nonwhite.sum(axis=0) > 0.55 * height)
        )
        if bounds[1] - bounds[0] > 0.10 * width
    ]
    if len(row_ranges) != 4 or len(column_ranges) != 5:
        raise ValueError(
            "Nitel figür panel düzeni algılanamadı: "
            f"{len(row_ranges)} satır, {len(column_ranges)} sütun"
        )
    return column_ranges, row_ranges


def build_qualitative_examples(
    source_path: Path,
    output_dir: Path,
) -> tuple[tuple[str, Path], ...]:
    output_dir.mkdir(parents=True, exist_ok=True)
    panel_labels = ("Input + all GT bbox", "Reference union", "SAM1", "SAM2", "SAM3")
    panel_positions = (
        (75, 80),
        (660, 80),
        (1245, 80),
        (367, 650),
        (952, 650),
    )
    label_font = _load_font(27, bold=True)
    legend_font = _load_font(21)
    legend = (
        "Yeşil: TP | Turuncu: FP | Pembe: FN | "
        "Her GT bbox ayrı istem, maskeler birleşik görünüm"
    )
    outputs: list[tuple[str, Path]] = []

    with Image.open(source_path) as source:
        source = source.convert("RGB")
        column_ranges, row_ranges = _detect_qualitative_panel_ranges(source)
        for (stratum, title), (y0, y1) in zip(
            QUALITATIVE_LABELS,
            row_ranges,
            strict=True,
        ):
            canvas = Image.new("RGB", (1800, 1170), "white")
            draw = ImageDraw.Draw(canvas)

            for label, (x0, x1), (x, y) in zip(
                panel_labels,
                column_ranges,
                panel_positions,
                strict=True,
            ):
                panel = source.crop((x0, y0, x1, y1))
                panel = panel.resize((480, 480), Image.Resampling.LANCZOS)
                label_box = draw.textbbox((0, 0), label, font=label_font)
                draw.text(
                    (
                        x + (480 - (label_box[2] - label_box[0])) / 2,
                        y - 42,
                    ),
                    label,
                    fill="black",
                    font=label_font,
                )
                canvas.paste(panel, (x, y))

            legend_box = draw.textbbox((0, 0), legend, font=legend_font)
            draw.text(
                (
                    (1800 - (legend_box[2] - legend_box[0])) / 2,
                    1133,
                ),
                legend,
                fill="black",
                font=legend_font,
            )
            output_path = output_dir / f"{stratum}.png"
            canvas.save(output_path, optimize=True)
            outputs.append((title, output_path))
    return tuple(outputs)


def build_docx(
    output_path: Path,
    *,
    spec: ReportSpec,
    detector_table: pd.DataFrame,
    tables: dict[tuple[str, str], pd.DataFrame],
    comparison_table: pd.DataFrame | None,
    comparison_note: str | None,
    qualitative_examples: tuple[tuple[str, Path], ...],
) -> None:
    detector_seeds = tuple(detector_table.attrs["detector_seeds"])
    document = Document()
    _configure_docx(document)

    title = document.add_heading(spec.title, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_heading("Scope", level=1)
    _add_docx_bullets(document, spec.scope_bullets)
    document.add_heading("Metric Logic", level=1)
    _add_docx_bullets(
        document,
        metric_bullets(spec.target_label, detector_seeds),
    )
    document.add_heading("Dataset Context", level=1)
    _add_docx_bullets(document, spec.context_bullets)

    document.add_heading("YOLO Detector BBox Metrics", level=1)
    _add_docx_bullets(document, detector_bullets(detector_seeds))
    _add_docx_table(document, detector_table)

    document.add_heading("Segmentation Tables", level=1)
    document.add_paragraph(
        "Başarı metrikleri 0 ile 1 arasındadır ve kırmızı-sarı-yeşil "
        "ölçeğiyle renklendirilmiştir."
    )
    for section in spec.reference_sections:
        document.add_heading(section.title, level=1)
        document.add_paragraph(section.note)
        for stratum, stratum_label in STRATA:
            frame = tables[(section.reference_type, stratum)]
            _add_docx_table(
                document,
                frame,
                title=stratum_label,
                note=segmentation_table_note(
                    section=section,
                    stratum=stratum,
                    frame=frame,
                    target_label=spec.target_label,
                ),
            )

    if comparison_table is not None and comparison_note is not None:
        document.add_heading("Reference Bias Comparison", level=1)
        document.add_paragraph(comparison_note)
        _add_docx_table(document, comparison_table)

    document.add_heading("Qualitative Examples", level=1)
    document.add_paragraph(
        "Her sayfa bir overlap × mask-area grubundan tek görüntüyü gösterir. "
        f"Görüntüdeki bütün GT {spec.target_label} kutuları modele ayrı istemler olarak "
        "verilmiş, üretilen instance maskeleri yalnız görsel sunum için "
        "birleştirilmiştir. Tablolar instance-level kalır. Yeşil TP, turuncu "
        "FP ve pembe FN pikselleridir."
    )
    for image_title, image_path in qualitative_examples:
        document.add_heading(image_title, level=2)
        document.add_picture(str(image_path), width=Inches(9.8))

    document.add_heading("Discussion", level=1)
    _add_docx_bullets(document, spec.discussion_bullets)
    document.save(output_path)


def _wrap_lines(text: str, width: int = 145) -> str:
    return "\n".join(textwrap.wrap(text, width=width))


def _add_pdf_text_pages(
    pdf: PdfPages,
    *,
    title: str,
    sections: Iterable[tuple[str, Iterable[str]]],
) -> None:
    fig = plt.figure(figsize=(11.69, 8.27))
    fig.patch.set_facecolor("white")
    plt.axis("off")
    y = 0.95
    fig.text(0.05, y, title, fontsize=20, weight="bold", va="top")
    y -= 0.08

    for section_title, bullets in sections:
        if y < 0.15:
            pdf.savefig(fig, bbox_inches="tight")
            plt.close(fig)
            fig = plt.figure(figsize=(11.69, 8.27))
            fig.patch.set_facecolor("white")
            plt.axis("off")
            y = 0.95
        fig.text(0.05, y, section_title, fontsize=13, weight="bold", va="top")
        y -= 0.035
        for bullet in bullets:
            wrapped = _wrap_lines(f"• {bullet}")
            line_count = wrapped.count("\n") + 1
            required = 0.028 * line_count
            if y - required < 0.05:
                pdf.savefig(fig, bbox_inches="tight")
                plt.close(fig)
                fig = plt.figure(figsize=(11.69, 8.27))
                fig.patch.set_facecolor("white")
                plt.axis("off")
                y = 0.95
            fig.text(0.07, y, wrapped, fontsize=8.2, va="top")
            y -= required
        y -= 0.015

    pdf.savefig(fig, bbox_inches="tight")
    plt.close(fig)


def _add_pdf_table(
    pdf: PdfPages,
    frame: pd.DataFrame,
    *,
    title: str,
    note: str | None = None,
) -> None:
    fig, axis = plt.subplots(figsize=(17.0, 9.5))
    fig.patch.set_facecolor("white")
    axis.axis("off")
    axis.set_title(title, fontsize=16, weight="bold", pad=35 if note else 18)
    if note:
        fig.text(
            0.5,
            0.91,
            _wrap_lines(note, width=190),
            ha="center",
            va="top",
            fontsize=8.2,
        )
    table = axis.table(
        cellText=frame.values,
        colLabels=frame.columns,
        cellLoc="center",
        loc="center",
        colColours=["#1F4E79"] * len(frame.columns),
    )
    table.auto_set_font_size(False)
    table.set_fontsize(6.2)
    table.scale(1.0, 1.55)
    for (row_index, column_index), cell in table.get_celld().items():
        cell.set_edgecolor("#666666")
        cell.set_linewidth(0.35)
        if row_index == 0:
            cell.set_text_props(color="white", weight="bold")
            continue
        column = frame.columns[column_index]
        if column in SUCCESS_COLUMNS:
            numeric = numeric_value(frame.iloc[row_index - 1, column_index])
            if numeric is not None:
                cell.set_facecolor("#" + interpolate_color(numeric))
    pdf.savefig(fig, bbox_inches="tight")
    plt.close(fig)


def _add_pdf_image_page(pdf: PdfPages, *, title: str, image_path: Path) -> None:
    with Image.open(image_path) as image:
        width, height = image.size
        fig_width = 11.69
        fig_height = max(8.27, fig_width * height / width + 0.6)
        fig, axis = plt.subplots(figsize=(fig_width, fig_height))
        fig.patch.set_facecolor("white")
        axis.axis("off")
        axis.set_title(title, fontsize=15, weight="bold", pad=12)
        axis.imshow(image)
        pdf.savefig(fig, bbox_inches="tight")
        plt.close(fig)


def build_pdf(
    output_path: Path,
    *,
    spec: ReportSpec,
    detector_table: pd.DataFrame,
    tables: dict[tuple[str, str], pd.DataFrame],
    comparison_table: pd.DataFrame | None,
    comparison_note: str | None,
    qualitative_examples: tuple[tuple[str, Path], ...],
) -> None:
    detector_seeds = tuple(detector_table.attrs["detector_seeds"])
    with PdfPages(output_path) as pdf:
        _add_pdf_text_pages(
            pdf,
            title=spec.title,
            sections=(
                ("Scope", spec.scope_bullets),
                (
                    "Metric Logic",
                    metric_bullets(spec.target_label, detector_seeds),
                ),
                ("Dataset Context", spec.context_bullets),
            ),
        )
        _add_pdf_table(
            pdf,
            detector_table,
            title="YOLO Detector BBox Metrics",
            note=(
                "Not: Bu tablo yalnız YOLO detector bbox başarısını gösterir; "
                f"maske metrikleri değildir. Değerler {_seed_note(detector_seeds)}."
            ),
        )
        for section in spec.reference_sections:
            for stratum, stratum_label in STRATA:
                frame = tables[(section.reference_type, stratum)]
                _add_pdf_table(
                    pdf,
                    frame,
                    title=stratum_label,
                    note=segmentation_table_note(
                        section=section,
                        stratum=stratum,
                        frame=frame,
                        target_label=spec.target_label,
                    ),
                )
        if comparison_table is not None and comparison_note is not None:
            _add_pdf_table(
                pdf,
                comparison_table,
                title="Reference Bias Comparison",
                note=comparison_note,
            )
        for image_title, image_path in qualitative_examples:
            _add_pdf_image_page(pdf, title=image_title, image_path=image_path)
        _add_pdf_text_pages(
            pdf,
            title="Discussion",
            sections=(("Findings", spec.discussion_bullets),),
        )


def _markdown_table(frame: pd.DataFrame) -> str:
    def clean(value: object) -> str:
        return str(value).replace("|", "\\|").replace("\n", " ")

    header = "| " + " | ".join(clean(column) for column in frame.columns) + " |"
    divider = "| " + " | ".join("---" for _ in frame.columns) + " |"
    rows = [
        "| " + " | ".join(clean(value) for value in row) + " |"
        for row in frame.itertuples(index=False, name=None)
    ]
    return "\n".join([header, divider, *rows])


def build_markdown(
    output_path: Path,
    *,
    spec: ReportSpec,
    detector_table: pd.DataFrame,
    tables: dict[tuple[str, str], pd.DataFrame],
    comparison_table: pd.DataFrame | None,
    comparison_note: str | None,
    qualitative_examples: tuple[tuple[str, Path], ...],
) -> None:
    detector_seeds = tuple(detector_table.attrs["detector_seeds"])
    def relative_image(path: Path) -> str:
        return Path(os.path.relpath(path, output_path.parent)).as_posix()

    lines = [
        f"# {spec.title}",
        "",
        "## Scope",
        "",
        *[f"- {bullet}" for bullet in spec.scope_bullets],
        "",
        "## Metric Logic",
        "",
        *[
            f"- {bullet}"
            for bullet in metric_bullets(spec.target_label, detector_seeds)
        ],
        "",
        "## Dataset Context",
        "",
        *[f"- {bullet}" for bullet in spec.context_bullets],
        "",
        "## YOLO Detector BBox Metrics",
        "",
        *[f"- {bullet}" for bullet in detector_bullets(detector_seeds)],
        "",
        _markdown_table(detector_table),
        "",
    ]
    for section in spec.reference_sections:
        lines.extend([f"## {section.title}", "", section.note, ""])
        for stratum, stratum_label in STRATA:
            frame = tables[(section.reference_type, stratum)]
            lines.extend(
                [
                    f"### {stratum_label}",
                    "",
                    segmentation_table_note(
                        section=section,
                        stratum=stratum,
                        frame=frame,
                        target_label=spec.target_label,
                    ),
                    "",
                    _markdown_table(frame),
                    "",
                ]
            )
    if comparison_table is not None and comparison_note is not None:
        lines.extend(
            [
                "## Reference Bias Comparison",
                "",
                comparison_note,
                "",
                _markdown_table(comparison_table),
                "",
            ]
        )
    lines.extend(
        [
            "## Qualitative Examples",
            "",
            "Her sayfa bir gruptan tek görüntüyü gösterir. Görüntüdeki bütün "
            f"GT {spec.target_label} kutuları modele ayrı istemler olarak verilmiş ve "
            "instance maskeleri yalnız bu görsel için birleştirilmiştir. "
            "Tablolar instance-level kalır. Yeşil TP, turuncu FP ve pembe FN "
            "piksellerini gösterir.",
            "",
        ]
    )
    for image_title, image_path in qualitative_examples:
        lines.extend(
            [
                f"### {image_title}",
                "",
                f"![{image_title}]({relative_image(image_path)})",
                "",
            ]
        )
    lines.extend(
        [
            "## Discussion",
            "",
            *[f"- {bullet}" for bullet in spec.discussion_bullets],
            "",
        ]
    )
    output_path.write_text("\n".join(lines), encoding="utf-8")


def write_report(
    *,
    spec: ReportSpec,
    output_dir: Path,
    aggregates_path: Path,
    detector_summary_path: Path,
    comparison_table: pd.DataFrame | None = None,
    comparison_note: str | None = None,
    extra_input_paths: tuple[Path, ...] = (),
) -> dict[str, Path]:
    aggregates = pd.read_csv(aggregates_path)
    detector_summary = pd.read_csv(detector_summary_path)
    output_dir.mkdir(parents=True, exist_ok=True)
    tables_dir = output_dir / "tables"
    tables_dir.mkdir(parents=True, exist_ok=True)
    qualitative_examples = build_qualitative_examples(
        spec.qualitative_image,
        output_dir / "qualitative",
    )

    detector_table = build_detector_table(
        detector_summary, dataset_id=spec.dataset_id
    )
    tables: dict[tuple[str, str], pd.DataFrame] = {}
    for section in spec.reference_sections:
        for stratum, _ in STRATA:
            frame = build_segmentation_table(
                aggregates,
                dataset_id=spec.dataset_id,
                reference_type=section.reference_type,
                stratum=stratum,
            )
            tables[(section.reference_type, stratum)] = frame
            frame.to_csv(
                tables_dir / f"{section.reference_type}__{stratum}.csv",
                index=False,
            )
    detector_table.to_csv(tables_dir / "detector_summary.csv", index=False)
    if comparison_table is not None:
        comparison_table.to_csv(
            tables_dir / "reference_sensitivity.csv", index=False
        )

    markdown_path = output_dir / f"{spec.slug}_full_metric_document.md"
    docx_path = output_dir / f"{spec.slug}_full_metric_document_colored.docx"
    pdf_path = output_dir / f"{spec.slug}_full_metric_document_colored.pdf"
    build_markdown(
        markdown_path,
        spec=spec,
        detector_table=detector_table,
        tables=tables,
        comparison_table=comparison_table,
        comparison_note=comparison_note,
        qualitative_examples=qualitative_examples,
    )
    build_docx(
        docx_path,
        spec=spec,
        detector_table=detector_table,
        tables=tables,
        comparison_table=comparison_table,
        comparison_note=comparison_note,
        qualitative_examples=qualitative_examples,
    )
    build_pdf(
        pdf_path,
        spec=spec,
        detector_table=detector_table,
        tables=tables,
        comparison_table=comparison_table,
        comparison_note=comparison_note,
        qualitative_examples=qualitative_examples,
    )

    manifest_path = output_dir / "report_manifest.json"
    output_files = [
        markdown_path,
        docx_path,
        pdf_path,
        *sorted(tables_dir.glob("*.csv")),
        *[path for _, path in qualitative_examples],
    ]
    manifest = {
        "schema_version": 2,
        "status": "completed",
        "study_id": spec.study_id,
        "dataset_id": spec.dataset_id,
        "report_format": "legacy_samrs_full_metric_colored",
        "inputs": {
            portable_manifest_path(path): sha256_file(path)
            for path in (
                aggregates_path,
                detector_summary_path,
                spec.qualitative_image,
                *extra_input_paths,
            )
        },
        "outputs": {
            portable_manifest_path(path): sha256_file(path)
            for path in output_files
        },
        "strata": [stratum for stratum, _ in STRATA],
        "reference_sections": [
            section.reference_type for section in spec.reference_sections
        ],
    }
    manifest_path.write_text(
        json.dumps(manifest, indent=2, ensure_ascii=False) + "\n",
        encoding="utf-8",
    )
    return {
        "markdown": markdown_path,
        "docx": docx_path,
        "pdf": pdf_path,
        "manifest": manifest_path,
    }
