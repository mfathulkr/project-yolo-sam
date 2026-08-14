from __future__ import annotations

import hashlib
import json
import re
import textwrap
from pathlib import Path
from typing import Iterable

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from matplotlib.backends.backend_pdf import PdfPages

from .paths import (
    BBOX_SOURCES,
    DATASETS,
    MODELS,
    REFERENCES,
    REPO_ROOT,
    STRATA,
    ExperimentSource,
)


plt.rcParams["pdf.fonttype"] = 42
plt.rcParams["ps.fonttype"] = 42


STRATUM_LABELS = {
    "overall": "Overall",
    "no_overlap__low_mask_area": "No Overlap × Low Mask Area",
    "no_overlap__high_mask_area": "No Overlap × High Mask Area",
    "overlap__low_mask_area": "Overlap × Low Mask Area",
    "overlap__high_mask_area": "Overlap × High Mask Area",
}
SCORE_COLUMNS = {
    "İnsan",
    "Yayımlanmış SAMRS",
    "SAM1 pseudo",
    "Yeniden SAM1",
    "SAM2 pseudo",
    "SAM3 pseudo",
    "Temel IoU",
    "Referans IoU",
    "Öğretmen IoU",
    "Diğerleri Ort.",
    "Referans Anlaşması",
    "Maskeler Arası Ortalama IoU",
    "Kendi Etiketiyle IoU",
    "Diğer SAM Etiketleriyle Ortalama IoU",
    "İnsan Etiketiyle IoU",
    "Yayımlanmış Etiketle IoU",
}
REFERENCE_SHORT = {
    "human": "İnsan",
    "published_samrs_reference": "Yayımlanmış SAMRS",
    "reproduced_pseudo_sam1": "Yeniden SAM1",
    "pseudo_sam1": "SAM1 pseudo",
    "pseudo_sam2": "SAM2 pseudo",
    "pseudo_sam3": "SAM3 pseudo",
}


def interpolate_color(value: float) -> str:
    value = max(0.0, min(1.0, value))
    red, yellow, green = (248, 105, 107), (255, 235, 132), (99, 190, 123)
    if value <= 0.5:
        ratio, left, right = value / 0.5, red, yellow
    else:
        ratio, left, right = (value - 0.5) / 0.5, yellow, green
    rgb = tuple(
        round(left[index] + (right[index] - left[index]) * ratio)
        for index in range(3)
    )
    return "".join(f"{channel:02X}" for channel in rgb)


def numeric(value: object) -> float | None:
    match = re.match(r"^\s*([0-9]+(?:\.[0-9]+)?)", str(value))
    if match is None:
        return None
    number = float(match.group(1))
    return number if 0.0 <= number <= 1.0 else None


def markdown_table(frame: pd.DataFrame) -> str:
    def clean(value: object) -> str:
        return str(value).replace("|", "\\|").replace("\n", " ")

    header = "| " + " | ".join(clean(column) for column in frame.columns) + " |"
    divider = "| " + " | ".join("---" for _ in frame.columns) + " |"
    rows = [
        "| " + " | ".join(clean(value) for value in row) + " |"
        for row in frame.itertuples(index=False, name=None)
    ]
    return "\n".join((header, divider, *rows))


def matrix_table(
    source: ExperimentSource,
    aggregates: pd.DataFrame,
    *,
    bbox_source: str,
    stratum: str,
) -> pd.DataFrame:
    selected = aggregates[
        (aggregates["bbox_source"] == bbox_source)
        & (aggregates["stratum"] == stratum)
    ]
    pivot = (
        selected.pivot(index="model", columns="reference_type", values="mean_iou")
        .loc[list(MODELS), list(source.reference_types)]
    )
    output = pd.DataFrame({"Model": [model.upper() for model in MODELS]})
    for reference_type in source.reference_types:
        output[REFERENCE_SHORT[reference_type]] = pivot[reference_type].map(
            lambda value: f"{float(value):.3f}"
        ).tolist()
    return output


def own_reference_for(source: ExperimentSource, model: str) -> str:
    if model == "sam1":
        return (
            "pseudo_sam1"
            if source.dataset_family == "isaid"
            else "reproduced_pseudo_sam1"
        )
    return f"pseudo_{model}"


def readable_affinity_table(
    source: ExperimentSource,
    aggregates: pd.DataFrame,
    *,
    bbox_source: str = "yolo_bbox",
    strata: Iterable[str] = ("overall",),
    include_bbox: bool = False,
) -> pd.DataFrame:
    """Show the core comparison without exposing statistical implementation jargon."""
    strata = tuple(strata)
    rows: list[dict[str, object]] = []
    pseudo_references = [
        reference
        for reference in source.reference_types
        if reference not in {"human", "published_samrs_reference"}
    ]
    for stratum in strata:
        selected = aggregates[
            (aggregates["bbox_source"] == bbox_source)
            & (aggregates["stratum"] == stratum)
        ]
        for model in MODELS:
            model_rows = selected[selected["model"] == model].set_index(
                "reference_type"
            )
            own_reference = own_reference_for(source, model)
            other_references = [
                reference
                for reference in pseudo_references
                if reference != own_reference
            ]
            own_iou = float(model_rows.loc[own_reference, "mean_iou"])
            other_iou = float(
                model_rows.loc[other_references, "mean_iou"].astype(float).mean()
            )
            row: dict[str, object] = {
                "Sahne Grubu": STRATUM_LABELS[stratum],
                "Model": model.upper(),
                "Kendi Etiketiyle IoU": f"{own_iou:.3f}",
                "Diğer SAM Etiketleriyle Ortalama IoU": f"{other_iou:.3f}",
                "Ek IoU": f"{own_iou - other_iou:+.3f}",
            }
            if include_bbox:
                row = {
                    "BBox": "GT bbox" if bbox_source == "gt_bbox" else "YOLO bbox",
                    **row,
                }
            rows.append(row)
    output = pd.DataFrame(rows)
    if len(strata) == 1:
        output = output.drop(columns="Sahne Grubu")
    return output


def baseline_to_own_table(
    source: ExperimentSource,
    effects: pd.DataFrame,
    *,
    bbox_source: str = "yolo_bbox",
) -> pd.DataFrame:
    rows: list[dict[str, object]] = []
    baseline_label = (
        "İnsan Etiketiyle IoU"
        if source.dataset_family == "isaid"
        else "Yayımlanmış Etiketle IoU"
    )
    selected = effects[effects["bbox_source"] == bbox_source]
    for model in MODELS:
        own_reference = own_reference_for(source, model)
        row = selected[
            (selected["model"] == model)
            & (selected["comparison_reference"] == own_reference)
        ].iloc[0]
        rows.append(
            {
                "Model": model.upper(),
                baseline_label: f"{float(row['baseline_mean_iou']):.3f}",
                "Kendi Etiketiyle IoU": f"{float(row['comparison_mean_iou']):.3f}",
                "Puan Değişimi": f"{float(row['delta_iou']):+.3f}",
            }
        )
    return pd.DataFrame(rows)


def advantage_table(frame: pd.DataFrame) -> pd.DataFrame:
    output = frame.copy()
    output["Model"] = output["teacher"].str.upper()
    output["BBox"] = output["bbox_source"].map(
        {"gt_bbox": "GT bbox", "yolo_bbox": "YOLO bbox"}
    )
    output["Referans"] = output["reference_type"].map(REFERENCE_SHORT)
    output["Öğretmen IoU"] = output["teacher_score"].map(
        lambda value: f"{float(value):.3f}"
    )
    output["Diğerleri Ort."] = output["other_models_mean"].map(
        lambda value: f"{float(value):.3f}"
    )
    output["Öğretmen Avantajı"] = output["teacher_advantage"].map(
        lambda value: f"{float(value):+.3f}"
    )
    output["Identity Control"] = output["identity_control"].map(
        {True: "Evet", False: "Hayır"}
    )
    return output[
        [
            "Model",
            "BBox",
            "Referans",
            "Öğretmen IoU",
            "Diğerleri Ort.",
            "Öğretmen Avantajı",
            "Identity Control",
        ]
    ]


def agreement_table(frame: pd.DataFrame) -> pd.DataFrame:
    output = frame.copy()
    output["Referans A"] = output["reference_a"].map(REFERENCE_SHORT)
    output["Referans B"] = output["reference_b"].map(REFERENCE_SHORT)
    output["Maskeler Arası Ortalama IoU"] = output["mean_instance_iou"].map(
        lambda value: f"{float(value):.3f}"
    )
    output["Nesne Sayısı"] = output["instance_count"].astype(int)
    return output[
        ["Referans A", "Referans B", "Maskeler Arası Ortalama IoU", "Nesne Sayısı"]
    ]


def empty_table(frame: pd.DataFrame) -> pd.DataFrame:
    output = frame.copy()
    output["Referans"] = output["reference_type"].map(REFERENCE_SHORT)
    output["Boş Maske"] = output["empty_count"].astype(int)
    output["Boş Oranı"] = output["empty_rate"].map(
        lambda value: f"{float(value):.3f}"
    )
    return output[["Referans", "Boş Maske", "Boş Oranı"]]


def ranking_table(frame: pd.DataFrame) -> pd.DataFrame:
    output = frame.copy()
    output["BBox"] = output["bbox_source"].map(
        {"gt_bbox": "GT bbox", "yolo_bbox": "YOLO bbox"}
    )
    output["Referans"] = output["reference_type"].map(REFERENCE_SHORT)
    output["Sıralama"] = output["ranking"]
    return output[["BBox", "Referans", "Sıralama"]]


def _shade(cell, color: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    properties.append(shading)


def _docx_table(
    document: Document,
    frame: pd.DataFrame,
    title: str,
    note: str | None = None,
) -> None:
    document.add_heading(title, level=2)
    if note:
        paragraph = document.add_paragraph(note)
        paragraph.paragraph_format.space_after = Pt(4)
        for run in paragraph.runs:
            run.font.name = "DejaVu Sans"
            run.font.size = Pt(8)
    table = document.add_table(rows=1, cols=len(frame.columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, column in enumerate(frame.columns):
        cell = table.rows[0].cells[index]
        cell.text = str(column)
        _shade(cell, "1F4E79")
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True
            run.font.size = Pt(6.5)
    for _, row in frame.iterrows():
        cells = table.add_row().cells
        for index, column in enumerate(frame.columns):
            value = row[column]
            cells[index].text = str(value)
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cells[index].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = "DejaVu Sans"
                    run.font.size = Pt(6.2)
            if column in SCORE_COLUMNS:
                score = numeric(value)
                if score is not None:
                    _shade(cells[index], interpolate_color(score))


def _docx_bullets(document: Document, bullets: Iterable[str]) -> None:
    for bullet in bullets:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(str(bullet))
        run.font.name = "DejaVu Sans"
        run.font.size = Pt(9)


def _configure_docx(document: Document) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = section.bottom_margin = Inches(0.42)
    section.left_margin = section.right_margin = Inches(0.48)
    for style_name in ("Normal", "Title", "Heading 1", "Heading 2"):
        document.styles[style_name].font.name = "DejaVu Sans"


def _pdf_text_page(pdf: PdfPages, title: str, bullets: Iterable[str]) -> None:
    figure = plt.figure(figsize=(11.69, 8.27))
    figure.text(0.055, 0.94, title, fontsize=17, fontweight="bold", va="top")
    y = 0.85
    for bullet in bullets:
        wrapped = textwrap.wrap(str(bullet), width=115)
        figure.text(0.07, y, "• " + wrapped[0], fontsize=9.3, va="top")
        y -= 0.034
        for line in wrapped[1:]:
            figure.text(0.09, y, line, fontsize=9.3, va="top")
            y -= 0.032
        y -= 0.014
    plt.axis("off")
    pdf.savefig(figure, bbox_inches="tight")
    plt.close(figure)


def _pdf_table_page(
    pdf: PdfPages,
    title: str,
    frame: pd.DataFrame,
    note: str | None = None,
) -> None:
    figure, axis = plt.subplots(figsize=(13.8, 8.0))
    axis.axis("off")
    figure.text(
        0.5,
        0.965,
        title,
        fontsize=15,
        fontweight="bold",
        ha="center",
        va="top",
    )
    if note:
        figure.text(0.5, 0.925, note, fontsize=8.5, ha="center", va="top")
    table_height = min(0.84, 0.055 * (len(frame) + 1))
    table_y = max(0.025, (0.90 - table_height) / 2)
    table = axis.table(
        cellText=frame.astype(str).values,
        colLabels=list(frame.columns),
        cellLoc="center",
        bbox=(0.0, table_y, 1.0, table_height),
    )
    table.auto_set_font_size(False)
    table.set_fontsize(
        6.2 if len(frame) > 30 else (6.8 if len(frame.columns) > 5 else 7.5)
    )
    for (row_index, column_index), cell in table.get_celld().items():
        if row_index == 0:
            cell.set_facecolor("#1F4E79")
            cell.get_text().set_color("white")
            cell.get_text().set_fontweight("bold")
            continue
        column = str(frame.columns[column_index])
        if column in SCORE_COLUMNS:
            score = numeric(frame.iloc[row_index - 1, column_index])
            if score is not None:
                cell.set_facecolor("#" + interpolate_color(score))
    pdf.savefig(figure, bbox_inches="tight")
    plt.close(figure)


def _write_manifest(path: Path, inputs: Iterable[Path], outputs: Iterable[Path]) -> None:
    def record(item: Path) -> dict[str, object]:
        return {
            "path": item.resolve().relative_to(REPO_ROOT).as_posix(),
            "bytes": item.stat().st_size,
            "sha256": hashlib.sha256(item.read_bytes()).hexdigest(),
        }

    path.write_text(
        json.dumps(
            {
                "schema_version": 1,
                "status": "completed",
                "inputs": [record(item) for item in inputs],
                "outputs": [record(item) for item in outputs],
            },
            ensure_ascii=False,
            indent=2,
        )
        + "\n",
        encoding="utf-8",
    )


def experiment_summary(
    source: ExperimentSource,
    aggregates: pd.DataFrame,
    affinity: pd.DataFrame,
) -> list[str]:
    baseline = source.reference_types[0]
    yolo = aggregates[
        (aggregates["stratum"] == "overall")
        & (aggregates["bbox_source"] == "yolo_bbox")
    ]
    baseline_rows = yolo[yolo["reference_type"] == baseline].sort_values(
        "mean_iou", ascending=False
    )
    best = baseline_rows.iloc[0]
    direct = affinity[
        (affinity["bbox_source"] == "yolo_bbox")
        & (affinity["stratum"] == "overall")
    ]
    direct_range = (
        float(direct["self_vs_cross_iou"].min()),
        float(direct["self_vs_cross_iou"].max()),
    )
    statements = [
        f"Temel referansta YOLO-bbox Overall Avg IoU bakımından en yüksek model {str(best['model']).upper()} ({float(best['mean_iou']):.3f}) olmuştur.",
        "Aynı dondurulmuş tahminler bütün referanslara karşı değerlendirildiği için sütunlar arasındaki fark yalnız referans maskesi değişiminin etkisini gösterir.",
        "GT bbox ile bir modelin kendi ürettiği maske yine aynı maskeye karşı ölçüldüğünde IoU'nun 1,000 olması beklenir; bu bağımsız başarı sonucu değildir.",
        (
            "YOLO bbox koşulunda her modelin kendi ürettiği etikette kazandığı ek IoU, "
            "aynı tahminlerin diğer iki SAM etiketindeki ortalaması çıkarılarak hesaplandı; "
            f"Overall aralık {direct_range[0]:+.3f} ile "
            f"{direct_range[1]:+.3f} arasındadır."
        ),
        "Tabaka sonuçları, etkinin kalabalık/örtüşen sahne ve hedef alanı koşullarında tutarlı olup olmadığını kontrol eder.",
    ]
    if source.dataset_family == "isaid":
        statements.append(
            "İnsan referansı bağımsız kontroldür; pseudo referans kaynaklı sıralama veya skor değişimi buna göre yorumlanır."
        )
    else:
        statements.append(
            "Yayımlanmış SAMRS referansı insan GT değildir. Bu nedenle sonuç, mutlak kalite karşılaştırmasından çok SAM-türevi referans yakınlığı analizidir."
        )
    return statements


def write_experiment_report(source: ExperimentSource) -> dict[str, Path]:
    analysis = source.analysis_root
    aggregates = pd.read_csv(analysis / "aggregate_metrics.csv")
    effects = pd.read_csv(analysis / "paired_reference_effects.csv")
    affinity = pd.read_csv(analysis / "paired_teacher_affinity_contrasts.csv")
    rankings = pd.read_csv(analysis / "ranking_by_reference.csv")
    agreements = pd.read_csv(analysis / "reference_agreement.csv")
    empties = pd.read_csv(analysis / "reference_empty_stats.csv")
    matrices = {
        (stratum, bbox): matrix_table(
            source, aggregates, bbox_source=bbox, stratum=stratum
        )
        for stratum in STRATA
        for bbox in BBOX_SOURCES
    }
    own_score_display = readable_affinity_table(source, aggregates)
    strata_score_display = readable_affinity_table(
        source,
        aggregates,
        strata=STRATA[1:],
    )
    baseline_own_display = baseline_to_own_table(source, effects)
    rankings_display = ranking_table(rankings)
    agreements_display = agreement_table(agreements)
    empties_display = empty_table(empties)
    summary = experiment_summary(source, aggregates, affinity)
    definitions = [
        (
            "Kapsam: 512 görüntü, dört sahne grubunun her birinde 128 görüntü ve "
            f"toplam {f'{source.instance_count:,}'.replace(',', '.')} "
            f"{source.target_label} nesnesi."
        ),
        "Avg IoU her nesne için ayrı hesaplanır ve bütün nesneler eşit ağırlıkla ortalanır.",
        "Model–referans matrislerinde satır değerlendirilen modeli, sütun kullanılan referans maskeyi, hücre ise Avg IoU değerini gösterir.",
        "Kendi Etiketiyle IoU, örneğin SAM2 tahmininin SAM2 tarafından üretilen referans maskeye göre puanıdır.",
        "Diğer SAM Etiketleriyle Ortalama IoU, aynı tahminin diğer iki SAM modelinin ürettiği maskelere göre aldığı iki puanın ortalamasıdır.",
        "Ek IoU, bu iki değerin farkıdır. Pozitif değer, modelin kendi etiketine göre ölçüldüğünde daha yüksek puan aldığını gösterir.",
        "Güven aralıkları ve ayrıntılı istatistiksel kontroller analiz CSV'lerinde saklanır; okunabilirliği korumak için bu özet tablolara basılmaz.",
        "Bu karşılaştırmalar ilk sonuçlar görüldükten sonra geliştirilmiştir; önceden kaydedilmiş doğrulayıcı test değildir ve çoklu karşılaştırma düzeltmesi uygulanmamıştır.",
        "Aynı-model karşılaştırması aynı dondurulmuş checkpoint ile sınırlıdır; farklı eğitim seed'i/checkpoint'i veya model ailesi düzeyinde genelleme test edilmemiştir.",
        "Bilinen pozitif nesnedeki boş pseudo referans eksik etikettir ve 0 puanlanır.",
    ]
    output_dir = source.reports_root / "cross_analysis"
    output_dir.mkdir(parents=True, exist_ok=True)
    slug = f"{source.experiment_id}_cross_reference_analysis"
    markdown_path = output_dir / f"{slug}.md"
    docx_path = output_dir / f"{slug}_colored.docx"
    pdf_path = output_dir / f"{slug}_colored.pdf"
    own_score_note = (
        "Ek IoU = Kendi Etiketiyle IoU − Diğer SAM Etiketleriyle Ortalama IoU. "
        "Pozitif değer, modelin kendi etiketinde daha yüksek puan aldığını gösterir."
    )
    report_tables = (
        (
            "Model Kendi Etiketiyle Ne Kadar Ek Puan Alıyor? · Overall · YOLO bbox",
            own_score_display,
            own_score_note,
        ),
        (
            "Sahne Gruplarına Göre Kendi Etiketindeki Ek Puan · YOLO bbox",
            strata_score_display,
            own_score_note,
        ),
        (
            "Temel Referanstan Kendi Etiketine Geçince Puan Değişimi · YOLO bbox",
            baseline_own_display,
            "Aynı tahmin sabit tutulur; yalnız puanın hesaplandığı referans maske değişir.",
        ),
        (
            "Referans Maskeler Birbirine Ne Kadar Benziyor?",
            agreements_display,
            "Bu tablo model başarısını değil, iki referans maske kümesinin birbirine benzerliğini gösterir.",
        ),
        ("Referansa Göre Model Sırası", rankings_display, None),
        ("Boş Üretilen Referans Maskeler", empties_display, None),
    )

    markdown = [
        f"# {source.display_name} Cross-Reference Analysis",
        "",
        "## Teknik Özet",
        "",
        *[f"- {item}" for item in summary],
        "",
        "## Kapsam ve Tanımlar",
        "",
        *[f"- {item}" for item in definitions],
        "",
    ]
    for stratum in STRATA:
        markdown.extend([f"## {STRATUM_LABELS[stratum]}", ""])
        for bbox in BBOX_SOURCES:
            markdown.extend(
                [
                    f"### {'GT bbox' if bbox == 'gt_bbox' else 'YOLO bbox'} Avg IoU",
                    "",
                    markdown_table(matrices[(stratum, bbox)]),
                    "",
                ]
            )
    for title, frame, note in report_tables:
        markdown.extend(
            [f"## {title}", "", *([note, ""] if note else []), markdown_table(frame), ""]
        )
    markdown_path.write_text("\n".join(markdown), encoding="utf-8")

    document = Document()
    _configure_docx(document)
    document.add_heading(
        f"{source.display_name} Cross-Reference Analysis",
        0,
    )
    document.add_heading("Teknik Özet", level=1)
    _docx_bullets(document, summary)
    document.add_heading("Kapsam ve Tanımlar", level=1)
    _docx_bullets(document, definitions)
    document.add_picture(
        str(source.figures_root / "model_reference_iou_matrix.png"),
        width=Inches(9.6),
    )
    for stratum in STRATA:
        for bbox in BBOX_SOURCES:
            _docx_table(
                document,
                matrices[(stratum, bbox)],
                f"{STRATUM_LABELS[stratum]} · {'GT bbox' if bbox == 'gt_bbox' else 'YOLO bbox'} Avg IoU",
                "Satır modeli, sütun referans maskeyi, hücre Avg IoU değerini gösterir.",
            )
    for title, frame, note in report_tables:
        _docx_table(document, frame, title, note)
    document.save(docx_path)

    with PdfPages(pdf_path) as pdf:
        _pdf_text_page(pdf, "Teknik Özet", summary)
        _pdf_text_page(pdf, "Kapsam ve Tanımlar", definitions)
        for stratum in STRATA:
            for bbox in BBOX_SOURCES:
                _pdf_table_page(
                    pdf,
                    f"{STRATUM_LABELS[stratum]} · {'GT bbox' if bbox == 'gt_bbox' else 'YOLO bbox'} Avg IoU",
                    matrices[(stratum, bbox)],
                    "Satır modeli, sütun referans maskeyi, hücre Avg IoU değerini gösterir.",
                )
        for title, frame, note in report_tables:
            _pdf_table_page(pdf, title, frame, note)
    outputs = (markdown_path, docx_path, pdf_path)
    inputs = tuple(analysis / name for name in (
        "aggregate_metrics.csv",
        "paired_teacher_affinity_contrasts.csv",
        "paired_reference_effects.csv",
        "ranking_by_reference.csv",
        "reference_agreement.csv",
        "reference_empty_stats.csv",
    ))
    _write_manifest(output_dir / "report_manifest.json", inputs, outputs)
    return {"markdown": markdown_path, "docx": docx_path, "pdf": pdf_path}


def main_control_table() -> pd.DataFrame:
    rows: list[dict[str, object]] = []
    for experiment_id, source in DATASETS.items():
        aggregates = pd.read_csv(source.analysis_root / "aggregate_metrics.csv")
        baseline = source.reference_types[0]
        overall = aggregates[
            (aggregates["stratum"] == "overall")
            & (aggregates["bbox_source"] == "yolo_bbox")
            & (aggregates["reference_type"] == baseline)
        ]
        for _, row in overall.iterrows():
            rows.append(
                {
                    "Deney": source.display_name,
                    "Temel Referans": REFERENCE_SHORT[baseline],
                    "Model": str(row["model"]).upper(),
                    "BBox": "YOLO bbox",
                    "Avg IoU": f"{float(row['mean_iou']):.3f}",
                    "Nesne Sayısı": int(row["instance_count"]),
                }
            )
    return pd.DataFrame(rows)


def main_affinity_table(experiment_ids: Iterable[str]) -> pd.DataFrame:
    frames: list[pd.DataFrame] = []
    for experiment_id in experiment_ids:
        source = DATASETS[experiment_id]
        aggregates = pd.read_csv(source.analysis_root / "aggregate_metrics.csv")
        frame = readable_affinity_table(source, aggregates)
        frame.insert(0, "Deney", source.display_name)
        frames.append(frame)
    return pd.concat(frames, ignore_index=True)


def main_reference_change_table(experiment_ids: Iterable[str]) -> pd.DataFrame:
    frames: list[pd.DataFrame] = []
    for experiment_id in experiment_ids:
        source = DATASETS[experiment_id]
        effects = pd.read_csv(source.analysis_root / "paired_reference_effects.csv")
        frame = baseline_to_own_table(source, effects)
        frame.insert(0, "Deney", source.display_name)
        frames.append(frame)
    return pd.concat(frames, ignore_index=True)


def samrs_affinity_table() -> pd.DataFrame:
    rows: list[dict[str, object]] = []
    for experiment_id in ("samrs_plane", "samrs_small_vehicle"):
        source = DATASETS[experiment_id]
        agreement = pd.read_csv(source.analysis_root / "reference_agreement.csv")
        selected = agreement[
            (agreement["reference_a"] == "published_samrs_reference")
            & (agreement["reference_b"] == "reproduced_pseudo_sam1")
        ]
        if len(selected) != 1:
            raise ValueError(f"SAMRS yayımlanmış/SAM1 anlaşması eksik: {experiment_id}")
        row = selected.iloc[0]
        rows.append(
            {
                "Deney": source.display_name,
                "Referans Çifti": "Yayımlanmış SAMRS ↔ yeniden SAM1",
                "Referans Anlaşması": f"{float(row['mean_instance_iou']):.3f}",
                "Nesne Sayısı": int(row["instance_count"]),
            }
        )
    return pd.DataFrame(rows)


def write_main_report(output_dir: Path) -> dict[str, Path]:
    output_dir.mkdir(parents=True, exist_ok=True)
    control = main_control_table()
    isaid_affinity = main_affinity_table(("isaid_plane", "isaid_small_vehicle"))
    isaid_change = main_reference_change_table(
        ("isaid_plane", "isaid_small_vehicle")
    )
    samrs_affinity = main_affinity_table(("samrs_plane", "samrs_small_vehicle"))
    samrs = samrs_affinity_table()
    summary = [
        "Dört deney aynı 512 görüntü / dört eşit 128 görüntülük sahne grubu / seed 42 / 1024×1024 / SAM1-2-3 / GT ve YOLO bbox protokolünü kullanır.",
        "iSAID Plane ve Small Vehicle deneylerinde insan anotasyonu bağımsız kontrol referansıdır. Ana karşılaştırma, aynı dondurulmuş model tahmininin kendi ürettiği etikette aldığı IoU ile diğer iki SAM etiketinde aldığı ortalama IoU arasındaki farktır.",
        "SAMRS Plane ve Small Vehicle deneylerinde yayımlanmış etiket insan GT değildir; SAM tabanlı üretim hattından geldiği için bu iki deney destekleyici SAM1-benzeri referans yakınlığı analizi olarak yorumlanır.",
        "GT bbox ile bir modelin kendi ürettiği maske yine aynı maskeye karşı ölçüldüğünde IoU'nun 1,000 olması beklenen matematiksel sonuçtur. Ana değerlendirme bu doğrudan eşitliği kıran YOLO bbox koşuludur.",
        "Dört deney tek bir ortalamada birleştirilmez. iSAID ve SAMRS ayrı anotasyon ürünleridir ancak ikisi de DOTA kökenli görüntüler içerdiği ve test kümeleri kısmen örtüştüğü için bağımsız dört replikasyon olarak yorumlanmaz.",
        "Sonuçlar, bu dondurulmuş checkpoint'ler ve seçilmiş test kapsamı içinde, pseudo etiket üreticisiyle bağımlı test referansının skor ve model seçimini etkileyebildiğini destekler; pseudo etiketlemenin eğitimde yararsız olduğunu göstermez.",
    ]
    limitations = [
        "Çalışma iki DOTA kökenli anotasyon ürünü ve iki hedef sınıfla sınırlıdır; dört deney bağımsız veri replikasyonları değildir.",
        "512 görüntülük testlerin tamamı hedef-pozitif olarak seçilmiştir. Detector AP değerleri resmi, negatif görüntüler de içeren benchmark AP'si değil bu seçilmiş pozitif test kapsamının kontrol metriğidir.",
        "SAMRS için bağımsız insan instance maskesi bulunmadığından mutlak segmentasyon kalitesi iddiası kurulamaz.",
        "Pseudo referanslar insan/yayımlanmış anotasyon kutularından gelen GT bbox ile, YOLO aday maskeleri ise tahmin kutularıyla üretilmiştir. Bu nedenle ölçülen fark yalnız sınır stilini izole etmez; checkpoint kimliği, GT/YOLO kutu farkı, prompt hassasiyeti ve maske biçiminin ortak etkileşimidir. Deney tam otomatik pseudo-etiketleme hattı değildir.",
        "Ana kendi-etiketi karşılaştırmaları ilk sonuçlar görüldükten sonra geliştirilmiş destekleyici analizlerdir; önceden kaydedilmiş doğrulayıcı test değildir ve çoklu karşılaştırma düzeltmesi uygulanmamıştır.",
        "Aynı-üretici etkisi aynı dondurulmuş SAM1/2/3 checkpoint'leri için ölçülmüştür. Farklı seed/checkpoint veya model ailesi düzeyinde genelleme bu çalışmada test edilmemiştir.",
        "YOLO yanlış pozitifleri detector mAP/precision/recall tablosunda ölçülür; instance maske ortalamasına sahte bir GT örneği olarak eklenmez. Bu nedenle maske tabloları tam uçtan uca instance-segmentation AP'si değildir.",
        "Her detector tek hedef sınıflıdır; bu nedenle detector mAP değeri o tek sınıfın AP değerine eşittir.",
    ]
    markdown_path = output_dir / "main_cross_analysis.md"
    docx_path = output_dir / "main_cross_analysis_colored.docx"
    pdf_path = output_dir / "main_cross_analysis_colored.pdf"
    markdown_path.write_text(
        "\n".join(
            [
                "# Teacher-Reference Bias Main Cross Analysis",
                "",
                "## Ana Sonuç",
                "",
                *[f"- {item}" for item in summary],
                "",
                "## Dört Deneyde Temel Referans Sonuçları",
                "",
                markdown_table(control),
                "",
                "## iSAID: Model Kendi Etiketiyle Ne Kadar Ek Puan Alıyor?",
                "",
                "Ek IoU = Kendi Etiketiyle IoU − Diğer SAM Etiketleriyle Ortalama IoU.",
                "",
                markdown_table(isaid_affinity),
                "",
                "## iSAID: İnsan Etiketinden Kendi Etiketine Geçince Ne Değişiyor?",
                "",
                "Aynı tahmin sabit tutulur; yalnız puanın hesaplandığı referans maske değişir.",
                "",
                markdown_table(isaid_change),
                "",
                "## SAMRS: Model Kendi Etiketiyle Ne Kadar Ek Puan Alıyor?",
                "",
                "Ek IoU = Kendi Etiketiyle IoU − Diğer SAM Etiketleriyle Ortalama IoU.",
                "",
                markdown_table(samrs_affinity),
                "",
                "## SAMRS Yayımlanmış Referans Yakınlığı",
                "",
                markdown_table(samrs),
                "",
                "## Sınırlılıklar",
                "",
                *[f"- {item}" for item in limitations],
                "",
            ]
        ),
        encoding="utf-8",
    )
    document = Document()
    _configure_docx(document)
    document.add_heading("Teacher-Reference Bias Main Cross Analysis", 0)
    document.add_heading("Ana Sonuç", level=1)
    _docx_bullets(document, summary)
    _docx_table(document, control, "Dört Deneyde Temel Referans Sonuçları")
    _docx_table(
        document,
        isaid_affinity,
        "iSAID: Model Kendi Etiketiyle Ne Kadar Ek Puan Alıyor?",
        "Ek IoU = kendi etiketindeki IoU − diğer iki SAM etiketindeki ortalama IoU.",
    )
    _docx_table(
        document,
        isaid_change,
        "iSAID: İnsan Etiketinden Kendi Etiketine Geçince Ne Değişiyor?",
        "Aynı tahmin sabit tutulur; yalnız puanın hesaplandığı referans maske değişir.",
    )
    _docx_table(
        document,
        samrs_affinity,
        "SAMRS: Model Kendi Etiketiyle Ne Kadar Ek Puan Alıyor?",
        "Ek IoU = kendi etiketindeki IoU − diğer iki SAM etiketindeki ortalama IoU.",
    )
    _docx_table(document, samrs, "SAMRS Yayımlanmış Referans Yakınlığı")
    document.add_heading("Sınırlılıklar", level=1)
    _docx_bullets(document, limitations)
    document.save(docx_path)
    with PdfPages(pdf_path) as pdf:
        _pdf_text_page(pdf, "Ana Sonuç", summary)
        _pdf_table_page(pdf, "Dört Deneyde Temel Referans Sonuçları", control)
        _pdf_table_page(
            pdf,
            "iSAID: Model Kendi Etiketiyle Ne Kadar Ek Puan Alıyor?",
            isaid_affinity,
            "Ek IoU = kendi etiketindeki IoU − diğer iki SAM etiketindeki ortalama IoU.",
        )
        _pdf_table_page(
            pdf,
            "iSAID: İnsan Etiketinden Kendi Etiketine Geçince Ne Değişiyor?",
            isaid_change,
            "Aynı tahmin sabit tutulur; yalnız puanın hesaplandığı referans maske değişir.",
        )
        _pdf_table_page(
            pdf,
            "SAMRS: Model Kendi Etiketiyle Ne Kadar Ek Puan Alıyor?",
            samrs_affinity,
            "Ek IoU = kendi etiketindeki IoU − diğer iki SAM etiketindeki ortalama IoU.",
        )
        _pdf_table_page(pdf, "SAMRS Yayımlanmış Referans Yakınlığı", samrs)
        _pdf_text_page(pdf, "Sınırlılıklar", limitations)
    outputs = (markdown_path, docx_path, pdf_path)
    inputs = tuple(
        source.analysis_root / filename
        for source in DATASETS.values()
        for filename in (
            "aggregate_metrics.csv",
            "paired_teacher_affinity_contrasts.csv",
            "paired_reference_effects.csv",
            "reference_agreement.csv",
        )
    )
    _write_manifest(output_dir / "report_manifest.json", inputs, outputs)
    return {"markdown": markdown_path, "docx": docx_path, "pdf": pdf_path}
