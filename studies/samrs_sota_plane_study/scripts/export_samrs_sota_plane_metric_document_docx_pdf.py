from __future__ import annotations

import argparse
import sys
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from matplotlib.backends.backend_pdf import PdfPages

STUDY_ROOT = Path(__file__).resolve().parents[1]
REPO_ROOT = STUDY_ROOT.parents[1]
ROOT = REPO_ROOT
SCRIPTS = STUDY_ROOT / "scripts"
if str(SCRIPTS) not in sys.path:
    sys.path.insert(0, str(SCRIPTS))

from export_isaid_metric_document_docx_pdf import (  # noqa: E402
    add_bullets,
    add_docx_table,
    add_pdf_image_page,
    add_pdf_table,
    add_pdf_text_page,
)
from write_samrs_sota_plane_metric_document import (  # noqa: E402
    CONTEXT_BULLETS,
    DETECTOR_METRIC_NOTE,
    DOCUMENT_TITLE,
    METRIC_BULLETS,
    RAW_SUMMARIES_DIR_NAME,
    SCOPE_BULLETS,
    SEGMENTATION_TABLES_NOTE,
    TABLES,
    discussion_lines,
    display_detector_metrics,
    display_summary,
    table_proxy_note,
)


def parse_args() -> argparse.Namespace:
    base_dir = STUDY_ROOT / "reports"
    parser = argparse.ArgumentParser(description="SAMRS SOTA plane metrik raporunu DOCX ve PDF olarak disa aktar.")
    parser.add_argument("--output-dir", type=Path, default=base_dir)
    parser.add_argument("--tables-dir", type=Path, default=base_dir / "tables" / "full_metric_document")
    parser.add_argument(
        "--visualizations-dir",
        type=Path,
        default=STUDY_ROOT / "results" / "pipelines" / "visualizations",
    )
    parser.add_argument(
        "--detector-metrics",
        type=Path,
        default=(
            STUDY_ROOT
            / "results"
            / "pipelines"
            / "detector_metrics"
            / "yolo_detector_eval_metrics.csv"
        ),
    )
    return parser.parse_args()


def load_summaries(tables_dir: Path) -> dict[str, pd.DataFrame]:
    raw_tables_dir = tables_dir / RAW_SUMMARIES_DIR_NAME
    if not raw_tables_dir.exists():
        raw_tables_dir = tables_dir
    summaries: dict[str, pd.DataFrame] = {}
    for table_key, _, _ in TABLES:
        summaries[table_key] = pd.read_csv(raw_tables_dir / f"summary_{table_key}.csv")
    return summaries


def collect_qualitative_images(visualizations_dir: Path) -> list[tuple[str, Path]]:
    if not visualizations_dir.exists():
        return []
    selected: list[tuple[str, Path]] = []
    seen: set[str] = set()
    for image_path in sorted(visualizations_dir.glob("*.png")):
        stratum = image_path.name.split("__", 2)
        if len(stratum) < 2:
            continue
        key = "__".join(stratum[:2])
        if key in seen:
            continue
        seen.add(key)
        title = key.replace("__", " / ").replace("_", " ").title()
        selected.append((title, image_path))
    return selected


def build_docx(
    output_path: Path,
    summaries: dict[str, pd.DataFrame],
    detector_table: pd.DataFrame,
    qualitative_images: list[tuple[str, Path]],
) -> None:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    for attr in ["top_margin", "bottom_margin", "left_margin", "right_margin"]:
        setattr(section, attr, Inches(0.35))

    title = doc.add_heading(DOCUMENT_TITLE, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("Scope", level=1)
    add_bullets(doc, SCOPE_BULLETS)

    doc.add_heading("Metric Logic", level=1)
    add_bullets(doc, METRIC_BULLETS)

    doc.add_heading("Dataset Context", level=1)
    add_bullets(doc, CONTEXT_BULLETS)

    doc.add_heading("YOLO Detector BBox Metrics", level=1)
    paragraph = doc.add_paragraph()
    paragraph.add_run(DETECTOR_METRIC_NOTE).font.size = Pt(10)
    if not detector_table.empty:
        add_docx_table(doc, detector_table)

    doc.add_heading("Segmentation Tables", level=1)
    doc.add_paragraph(SEGMENTATION_TABLES_NOTE)
    for table_key, title_text, _ in TABLES:
        add_docx_table(
            doc,
            display_summary(summaries[table_key]),
            title=title_text,
            note=table_proxy_note(summaries[table_key]),
        )

    if qualitative_images:
        doc.add_heading("Qualitative Examples", level=1)
        for title_text, image_path in qualitative_images:
            doc.add_heading(title_text, level=2)
            doc.add_picture(str(image_path), width=Inches(9.8))

    doc.add_heading("Discussion", level=1)
    add_bullets(doc, discussion_lines(summaries))
    doc.save(output_path)


def build_pdf(
    output_path: Path,
    summaries: dict[str, pd.DataFrame],
    detector_table: pd.DataFrame,
    qualitative_images: list[tuple[str, Path]],
) -> None:
    with PdfPages(output_path) as pdf:
        add_pdf_text_page(
            pdf,
            DOCUMENT_TITLE,
            [
                ("Scope", SCOPE_BULLETS),
                ("Metric Logic", METRIC_BULLETS),
                ("Dataset Context", CONTEXT_BULLETS),
            ],
        )
        if not detector_table.empty:
            add_pdf_table(pdf, detector_table, "YOLO Detector BBox Metrics", note=DETECTOR_METRIC_NOTE)
        for table_key, title_text, _ in TABLES:
            add_pdf_table(
                pdf,
                display_summary(summaries[table_key]),
                title_text,
                note=table_proxy_note(summaries[table_key]),
            )
        for title_text, image_path in qualitative_images:
            add_pdf_image_page(pdf, title_text, image_path)
        add_pdf_text_page(pdf, "Discussion", [("Findings", discussion_lines(summaries))])


def main() -> None:
    args = parse_args()
    args.output_dir.mkdir(parents=True, exist_ok=True)
    summaries = load_summaries(args.tables_dir)
    detector_table = display_detector_metrics(args.detector_metrics)
    qualitative_images = collect_qualitative_images(args.visualizations_dir)

    docx_path = args.output_dir / "samrs_sota_plane_full_metric_document_colored.docx"
    pdf_path = args.output_dir / "samrs_sota_plane_full_metric_document_colored.pdf"
    build_docx(docx_path, summaries, detector_table, qualitative_images)
    build_pdf(pdf_path, summaries, detector_table, qualitative_images)
    print(docx_path)
    print(pdf_path)


if __name__ == "__main__":
    main()
