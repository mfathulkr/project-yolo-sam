from __future__ import annotations

import argparse
import json
import shutil
import sys
import textwrap
from datetime import date
from pathlib import Path
from typing import Any

import matplotlib

matplotlib.use("Agg")

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from pandas.errors import EmptyDataError
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from matplotlib.backends.backend_pdf import PdfPages
from matplotlib.colors import Normalize

STUDY_ROOT = Path(__file__).resolve().parents[1]
REPO_ROOT = STUDY_ROOT.parents[1]
ROOT = REPO_ROOT
for source_root in (STUDY_ROOT / "src", REPO_ROOT / "src"):
    if str(source_root) not in sys.path:
        sys.path.insert(0, str(source_root))

from teacher_reference_bias.reporting.analysis import sha256_file


TITLE = (
    "Öğretmen Kendi Referansına Karşı Ne Kadar İyi? "
    "Uzaktan Algılama Segmentasyonunda SAM Üretimli Test Maskelerinin "
    "Değerlendirme Yanlılığı"
)
SHORT_TITLE = "SAM Üretimli Referanslarda Değerlendirme Yanlılığı"
MODELS = ("sam1", "sam2", "sam3")
MODEL_LABELS = {"sam1": "SAM1", "sam2": "SAM2", "sam3": "SAM3"}
DATASET_LABELS = {
    "isaid_plane": "iSAID Plane",
    "samrs_sota_plane": "SAMRS SOTA Plane",
}
REFERENCE_LABELS = {
    "human": "İnsan",
    "pseudo_sam1": "SAM1 pseudo",
}
PARTIAL_WARNING = (
    "KISMİ TASLAK: detector ve YOLO-bbox üç-seed sonuçları bekleniyor."
)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Build the source-backed six-page teacher-reference bias paper."
    )
    parser.add_argument(
        "--study-root",
        type=Path,
        default=STUDY_ROOT / "results",
    )
    parser.add_argument(
        "--output-dir",
        type=Path,
        default=ROOT / "studies/teacher_reference_bias_v1/reports/paper",
    )
    parser.add_argument(
        "--allow-partial",
        action="store_true",
        help="Build a clearly marked draft before all detector seeds are complete.",
    )
    return parser.parse_args()


def read_csv(path: Path) -> pd.DataFrame:
    if not path.exists() or path.stat().st_size == 0:
        return pd.DataFrame()
    try:
        return pd.read_csv(path)
    except EmptyDataError:
        return pd.DataFrame()


def load_evidence(study_root: Path) -> dict[str, Any]:
    analysis = study_root / "analysis"
    shared = analysis / "shared_human_reference_audit"
    metadata_paths = {
        dataset_id: (
            STUDY_ROOT / "data" / "prepared" / dataset_id / "test" / "metadata.csv"
        )
        for dataset_id in DATASET_LABELS
    }
    strata_definitions: dict[str, dict[str, float | int]] = {}
    for dataset_id, metadata_path in metadata_paths.items():
        metadata = read_csv(metadata_path)
        if metadata.empty:
            continue
        strata_definitions[dataset_id] = {
            "images": int(len(metadata)),
            "area_threshold": float(metadata["area_threshold"].iloc[0]),
            "no_overlap_iou_max": float(
                metadata["no_overlap_iou_max"].iloc[0]
            ),
            "overlap_iou_min": float(
                metadata["overlap_iou_min"].iloc[0]
            ),
        }
    evidence: dict[str, Any] = {
        "aggregates": read_csv(analysis / "aggregate_metrics.csv"),
        "comparisons": read_csv(analysis / "paired_model_comparisons.csv"),
        "inflation": read_csv(analysis / "reference_inflation.csv"),
        "rankings": read_csv(analysis / "ranking_comparisons.csv"),
        "detector": read_csv(analysis / "detector_seed_summary.csv"),
        "segmentation": read_csv(analysis / "segmentation_seed_summary.csv"),
        "shared_summary": read_csv(shared / "model_dual_reference_summary.csv"),
        "shared_quality": read_csv(shared / "reference_quality_summary.csv"),
        "shared_unique_objects": read_csv(
            shared / "unique_human_object_sensitivity.csv"
        ),
        "shared_ci": json.loads(
            (shared / "model_reference_inflation_ci.json").read_text(
                encoding="utf-8"
            )
        ),
        "shared_ranking": json.loads(
            (shared / "ranking_comparison.json").read_text(encoding="utf-8")
        ),
        "strata_definitions": strata_definitions,
    }
    evidence["input_paths"] = [
        path
        for path in (
            analysis / "aggregate_metrics.csv",
            analysis / "paired_model_comparisons.csv",
            analysis / "reference_inflation.csv",
            analysis / "ranking_comparisons.csv",
            analysis / "detector_seed_summary.csv",
            analysis / "segmentation_seed_summary.csv",
            analysis / "manifest.json",
            shared / "model_dual_reference_summary.csv",
            shared / "reference_quality_summary.csv",
            shared / "model_reference_inflation_ci.json",
            shared / "ranking_comparison.json",
            shared / "unique_human_object_sensitivity.csv",
            shared / "manifest.json",
            *metadata_paths.values(),
        )
        if path.exists()
    ]
    return evidence


def completion_issues(evidence: dict[str, Any]) -> list[str]:
    issues: list[str] = []
    strata_definitions = evidence["strata_definitions"]
    if set(strata_definitions) != set(DATASET_LABELS) or any(
        int(row["images"]) != 128
        or float(row["no_overlap_iou_max"]) != 0.0
        or float(row["overlap_iou_min"]) != 0.001
        for row in strata_definitions.values()
    ):
        issues.append("Frozen image-level strata definitions are incomplete")
    detector = evidence["detector"]
    if detector.empty:
        issues.append("Detector seed summary is empty")
    else:
        for dataset_id in DATASET_LABELS:
            rows = detector[detector["dataset_id"] == dataset_id]
            if rows.empty or int(rows.iloc[0]["seed_count"]) < 3:
                issues.append(f"{dataset_id} does not contain three detector seeds")

    segmentation = evidence["segmentation"]
    expected = {
        ("isaid_plane", model, reference)
        for model in MODELS
        for reference in ("human", "pseudo_sam1")
    } | {
        ("samrs_sota_plane", model, "pseudo_sam1")
        for model in MODELS
    }
    observed: set[tuple[str, str, str]] = set()
    if not segmentation.empty and "seed_count" in segmentation:
        for _, row in segmentation.iterrows():
            key = (
                str(row["dataset_id"]),
                str(row["model"]),
                str(row["reference_type"]),
            )
            if int(row["seed_count"]) >= 3:
                observed.add(key)
    missing = sorted(expected - observed)
    if missing:
        issues.append(f"Missing three-seed YOLO-bbox summaries: {missing}")

    intervals = evidence["shared_ci"]
    if len(intervals) != 3 or any(
        int(row["bootstrap_samples"]) != 10_000 for row in intervals
    ):
        issues.append("Shared human audit is not the frozen 10,000-bootstrap result")
    return issues


def gt_bbox_table(aggregates: pd.DataFrame) -> pd.DataFrame:
    selected = aggregates[
        (aggregates["bbox_source"] == "gt_bbox")
        & (aggregates["stratum"] == "overall")
    ].copy()
    rows: list[dict[str, object]] = []
    conditions = (
        ("isaid_plane", "human"),
        ("isaid_plane", "pseudo_sam1"),
        ("samrs_sota_plane", "pseudo_sam1"),
    )
    for dataset_id, reference_type in conditions:
        for model in MODELS:
            match = selected[
                (selected["dataset_id"] == dataset_id)
                & (selected["reference_type"] == reference_type)
                & (selected["model"] == model)
            ]
            if match.empty:
                continue
            row = match.iloc[0]
            rows.append(
                {
                    "Veri / referans": (
                        f"{DATASET_LABELS[dataset_id]} / "
                        f"{REFERENCE_LABELS[reference_type]}"
                    ),
                    "Model": MODEL_LABELS[model],
                    "IoU": float(row["mean_iou"]),
                    "Dice": float(row["mean_dice"]),
                    "Precision": float(row["mean_precision"]),
                    "Recall": float(row["mean_recall"]),
                    "Boundary IoU": float(row["mean_boundary_iou"]),
                }
            )
    return pd.DataFrame(rows)


def shared_reference_table(
    summary: pd.DataFrame,
    intervals: list[dict[str, Any]],
) -> pd.DataFrame:
    overall = summary[summary["stratum"] == "overall"]
    score = overall.pivot(
        index="model",
        columns="reference_type",
        values="mean_iou",
    )
    interval_by_model = {str(row["model"]): row for row in intervals}
    rows = []
    for model in MODELS:
        interval = interval_by_model[model]
        rows.append(
            {
                "Model": MODEL_LABELS[model],
                "İnsan IoU": float(score.loc[model, "human"]),
                "SAM1 pseudo IoU": float(score.loc[model, "pseudo_sam1"]),
                "Enflasyon": float(interval["estimate"]),
                "%95 GA": (
                    f"[{float(interval['lower']):.3f}, "
                    f"{float(interval['upper']):.3f}]"
                ),
            }
        )
    return pd.DataFrame(rows)


def detector_table(detector: pd.DataFrame) -> pd.DataFrame:
    if detector.empty:
        return pd.DataFrame(
            columns=[
                "Veri",
                "Seed",
                "Eşik",
                "P@IoU50",
                "R@IoU50",
                "AP50",
                "AP75",
                "AP90",
                "AP50-95",
            ]
        )
    rows = []
    for _, row in detector.sort_values("dataset_id").iterrows():
        rows.append(
            {
                "Veri": DATASET_LABELS[str(row["dataset_id"])],
                "Seed": int(row["seed_count"]),
                "Eşik": float(row["fixed_confidence_threshold_mean"]),
                "P@IoU50": float(
                    row["precision_at_bbox_iou50_mean"]
                ),
                "R@IoU50": float(row["recall_at_bbox_iou50_mean"]),
                "AP50": float(row["bbox_AP50_mean"]),
                "AP75": float(row["bbox_AP75_mean"]),
                "AP90": float(row["bbox_AP90_mean"]),
                "AP50-95": float(row["bbox_AP50_95_mean"]),
            }
        )
    return pd.DataFrame(rows)


def yolo_segmentation_table(segmentation: pd.DataFrame) -> pd.DataFrame:
    if segmentation.empty or "seed_count" not in segmentation:
        return pd.DataFrame(
            columns=["Veri / referans", "Model", "Seed", "IoU", "IoU std"]
        )
    rows = []
    for _, row in segmentation.sort_values(
        ["dataset_id", "reference_type", "model"]
    ).iterrows():
        rows.append(
            {
                "Veri / referans": (
                    f"{DATASET_LABELS[str(row['dataset_id'])]} / "
                    f"{REFERENCE_LABELS[str(row['reference_type'])]}"
                ),
                "Model": MODEL_LABELS[str(row["model"])],
                "Seed": int(row["seed_count"]),
                "IoU": float(row["mean_iou_seed_mean"]),
                "IoU std": float(row["mean_iou_seed_std"]),
                "Dice": float(row["mean_dice_seed_mean"]),
            }
        )
    return pd.DataFrame(rows)


def markdown_table(frame: pd.DataFrame, decimals: int = 3) -> str:
    if frame.empty:
        return "_Henüz tamamlanmış sonuç yok._"
    display = frame.copy()
    for column in display.columns:
        if pd.api.types.is_float_dtype(display[column]):
            display[column] = display[column].map(lambda value: f"{value:.{decimals}f}")
    header = "| " + " | ".join(str(column) for column in display.columns) + " |"
    divider = "| " + " | ".join("---" for _ in display.columns) + " |"
    rows = [
        "| " + " | ".join(str(value) for value in row) + " |"
        for row in display.itertuples(index=False, name=None)
    ]
    return "\n".join([header, divider, *rows])


def build_markdown(
    evidence: dict[str, Any],
    *,
    issues: list[str],
    figure_names: dict[str, str],
) -> str:
    gt_table = gt_bbox_table(evidence["aggregates"])
    shared_table = shared_reference_table(
        evidence["shared_summary"],
        evidence["shared_ci"],
    )
    det_table = detector_table(evidence["detector"])
    yolo_table = yolo_segmentation_table(evidence["segmentation"])
    quality = evidence["shared_quality"]
    quality_overall = quality[quality["stratum"] == "overall"].iloc[0]
    unique_sensitivity = evidence["shared_unique_objects"].set_index("model")
    strata_definitions = evidence["strata_definitions"]
    isaid_area_threshold = (
        100.0 * float(strata_definitions["isaid_plane"]["area_threshold"])
    )
    samrs_area_threshold = (
        100.0
        * float(strata_definitions["samrs_sota_plane"]["area_threshold"])
    )
    status = (
        "TAMAMLANMIŞ FROZEN SONUÇ"
        if not issues
        else "KISMİ TASLAK - EKSİK: " + "; ".join(issues)
    )
    return f"""# {TITLE}

**Durum:** {status}
**Çalışma kimliği:** `teacher_reference_bias_v1`
**Tarih:** {date.today().isoformat()}

## Öz

Otomatik pseudo-maskeler büyük ölçekli segmentation pretraining için değerlidir;
ancak bu maskeler bağımsız test ground truth'u olarak kullanıldığında referansı
üreten modeli kayırabilir. Bu çalışma, söz konusu **öğretmen-referans
yakınlığı (teacher-reference affinity)** etkisini uzaktan algılama instance segmentation bağlamında
kontrollü olarak ölçmektedir. iSAID ve SAMRS SOTA üzerinde aynı `plane`
sınıfı, aynı 1024x1024 giriş, aynı 128 test görüntüsü, aynı dört
`overlap x mask area` katmanı ve aynı SAM1/SAM2/SAM3 bbox-prompted pipeline'ı
kullanılmıştır. SAMRS SOTA maskelerinin resmi SAM1 ViT-H ve RHBox üretim
provenance'ı exhaustive geometri denetimiyle doğrulanmıştır. Ayrıca 126 SAMRS
tile'ı aynı DOTA görüntülerindeki bağımsız iSAID insan annotation'larına piksel
düzeyinde eşlenmiş; 770 benzersiz uçağın 1.033 crop görünümü için aynı
tahminler iki referansa karşı ölçülmüştür. SAM1'in ortalama IoU'su insan referansında
`{shared_table.iloc[0]['İnsan IoU']:.3f}`, kendi pseudo referansında
`{shared_table.iloc[0]['SAM1 pseudo IoU']:.3f}` olmuş; skor enflasyonu
`{shared_table.iloc[0]['Enflasyon']:.3f}` olarak ölçülmüştür. Bulgular SAMRS'nin
weak supervision değerini reddetmez; öğretmen üretimli test maskelerinin
bağımsız benchmark referansı olarak yorumlanmasına sınır koyar.

**Anahtar kelimeler:** remote sensing, Segment Anything, pseudo-label,
evaluation bias, instance segmentation, teacher-reference affinity

## 1. Giriş

Pixel-level uzaktan algılama annotation'ı pahalıdır. SAMRS bu maliyeti azaltmak
için mevcut detection bbox'larını SAM1 prompt'u olarak kullanır ve 1,6
milyondan fazla instance maskesi üretir. Bu yaklaşım pretraining için ölçek
sağlar. Bununla birlikte, aynı öğretmen model kendi ürettiği maskelere karşı
değerlendirildiğinde yüksek skor gerçek nesne sınırlarına uyum ile öğretmenin
karar stilini yeniden üretme etkisini karıştırabilir.

Bu bildirinin katkıları:

1. iSAID ve SAMRS için kaynak-sahne güvenli, eşlenmiş değerlendirme protokolü.
2. Aynı tahminlerin insan ve SAM1 pseudo referansında eşleştirilmiş ölçümü.
3. Kaynak-sahne kümeli 10.000 bootstrap ile referans enflasyonu aralıkları.
4. SAM1, SAM2 ve SAM3 için GT-bbox ile YOLO-bbox sonuçlarının ayrılması.
5. Pseudo-label eğitim yararı ile pseudo-referans benchmark geçerliliğinin
   açık ayrımı.

## 2. İlgili Çalışmalar

SAMRS, DOTA, DIOR, FAIR1M ve HRSC2016 detection annotation'larını SAM ile
maskeye dönüştürür ve ana kullanımını segmentation pretraining olarak
temellendirir. Brachmann ve arkadaşları pseudo-ground-truth üreten referans
algoritmasının benzer re-localisation yöntemlerini kayırabildiğini göstermiştir.
Arazo ve arkadaşlarının confirmation bias çalışması ise hatalı pseudo-label'ın
eğitimde kendini pekiştirebildiğini açıklar. Bizim problemimiz eğitim dinamiği
değil, ölçüm geçerliliği problemidir. SOPSeg'in iSAID ablation'ı küçük
nesnelerde en büyük kazancın decoder'dan önce region-adaptive magnification ve
oriented prompt'tan geldiğini gösterir; bu nedenle veri hazırlama ve prompt
geometrisi model isminden ayrı kontrol edilmelidir.

## 3. Materyal ve Yöntem

### 3.1 Veri ve split

- Hedef sınıf iki veri setinde de `plane`.
- iSAID referansı insan tarafından bağımsız çizilmiş instance polygon'larıdır.
- Tile maskeleri kayıpsız COCO RLE olarak saklanmış; boş maske ve piksel alanı
  uyuşmazlığı exhaustive olarak sıfır doğrulanmıştır.
- iSAID GT prompt'u bu resmi insan annotation'ındaki eksen hizalı bbox'tır.
- SAMRS SOTA referansı SAM1 ViT-H ve original DOTA RHBox prompt'larıyla
  üretilmiş pseudo-maskedir.
- Test seti veri seti başına 128 görüntüdür.
- Dört strata'nın her birinde 32 görüntü bulunur.
- Strata görüntü/tile düzeyindedir: `No Overlap` maksimum plane bbox-pair
  IoU değerinin tam `0`, `Overlap` ise en az `0,001` olmasıdır; aradaki belirsiz
  görüntüler örnekleme havuzuna alınmaz.
- Low/high ayrımı görüntüdeki instance mask alanları toplamının görüntü alanına
  oranının veri seti içi medyanıyla yapılır. Dondurulmuş eşikler iSAID için
  `%{isaid_area_threshold:.3f}` (insan maskesi), SAMRS için
  `%{samrs_area_threshold:.3f}` (SAM1 pseudo maskesi) değeridir. Bu nedenle
  dört strata cross-dataset nedensel kanıt değil, betimleyici zorluk
  kırılımıdır.
- Train, validation ve test arasında source scene kesişimi sıfırdır.
- Hiçbir GT prompt SAM1 pseudo-maskeden yeniden türetilmez. iSAID insan bbox'ı
  ile SAMRS original detection RHBox provenance farkı nedeniyle cross-dataset
  GT-bbox mutlak skorları betimleyicidir.

### 3.2 Dual-reference tasarımı

Kontrollü iSAID deneyinde SAM1 GT-bbox tahmini ikinci pseudo referans
olarak dondurulmuştur. Ortak görüntü denetiminde SAMRS tile'ları iSAID
kaynaklarına geri eşlenmiş ve official SAMRS pseudo-mask ile iSAID insan maskesi
aynı instance üzerinde kullanılmıştır. Prediction sabit kalır; yalnız referans
değişir.

### 3.3 Metrikler ve istatistik

Ana metrik instance-level mask IoU'dur. Dice, pixel precision, pixel recall,
Boundary IoU ve Success@0.50/0.75/0.90 destekleyici metriklerdir. Detector için
COCO bbox AP50, AP75, AP90 ve AP50-95 ayrı raporlanır. Her seed'in YOLO
confidence eşiği validation setinde bbox IoU 0,50 için F1'i en yüksek yapan
noktada dondurulur; test eşik seçimine girmez. Confidence interval hesabında
bağımsız gözlem birimi tile değil source scene'dir. Pairwise model
karşılaştırmalarında source-scene ortalama farkları üzerinde Wilcoxon testi ve
Holm düzeltmesi kullanılır.

YOLO-bbox koşulunda bbox IoU >= 0,50 ile eşleşmeyen GT için boş maske yazılır
ve instance skoru sıfır olur. Eşleşmeyen detector tahminleri detector AP
hesabında false positive, ikincil image-union maskesinde ise tahmin olarak
korunur; bir GT instance satırına yapay biçimde atanmaz. Bu nedenle
instance-level mask tablosu COCO mask AP değildir.

## 4. Sonuçlar

### 4.1 GT-bbox sonuçları

{markdown_table(gt_table)}

### 4.2 Aynı 1.033 tile-instance tahmini, iki referans

SAMRS pseudo-mask ile iSAID insan maskesi arasındaki ortalama IoU
`{float(quality_overall['mean_iou']):.3f}`'tür. Bu değer pseudo referansın
yüksek kaliteli fakat insan ground truth ile özdeş olmadığını gösterir.
Bu `1.033` satır, SAMRS'nin örtüşen tile'ları nedeniyle `770` benzersiz
insan-anotasyonlu uçağın farklı crop görünümleridir; belirsizlik hesabı `35`
kaynak sahne düzeyinde kümelenmiştir. Her benzersiz uçağın görünümleri önce
kendi içinde ortalandığında IoU enflasyonu SAM1, SAM2 ve SAM3 için sırasıyla
`{float(unique_sensitivity.loc['sam1', 'mean_iou_inflation']):.3f}`,
`{float(unique_sensitivity.loc['sam2', 'mean_iou_inflation']):.3f}` ve
`{float(unique_sensitivity.loc['sam3', 'mean_iou_inflation']):.3f}` olarak
aynı yönde kalmıştır.

{markdown_table(shared_table)}

![Aynı tahminlerin iki referanstaki skoru]({figure_names['shared']})

İnsan ve pseudo referansta model sırası bu örneklemde değişmemiştir
(`Spearman = {float(evidence['shared_ranking']['spearman_correlation']):.1f}`,
`Kendall tau = {float(evidence['shared_ranking']['kendall_tau']):.1f}`).
Dolayısıyla ana bulgu ranking reversal değil, model ailesine göre farklılaşan
güçlü skor enflasyonudur.

### 4.3 YOLO detector

{markdown_table(det_table)}

### 4.4 YOLO-bbox segmentation

{markdown_table(yolo_table)}

![GT-bbox strata sonuçları]({figure_names['strata']})

## 5. Tartışma

SAM1'in official SAMRS referansında yaklaşık kusursuz görünmesi, bağımsız insan
sınırlarına göre kusursuz olduğu anlamına gelmez. Aynı 1.033 tahminde
referansın değiştirilmesi SAM1 için yaklaşık 0,35 IoU farkı üretmiştir. SAM2 ve
SAM3 de pseudo referansta yükselmiş, ancak artış öğretmen SAM1 için en büyük
olmuştur. Bu desen genel veri seti zorluğuyla açıklanamaz; görüntü, instance,
bbox ve tahmin sabittir.

Sonuç SAMRS'nin değersiz olduğunu göstermez. Pseudo-maskeler pretraining,
distillation ve weak supervision için yararlı olabilir. Geçerli downstream
fayda bağımsız insan etiketli test setinde ölçülmelidir. Pseudo-mask benchmark
olarak kullanılacaksa üretici model, checkpoint, prompt türü ve insan denetimi
subset açıkça raporlanmalıdır.

SOPSeg bağlamı, remote sensing small-object segmentation'da sonraki ilerlemenin
yalnız yeni SAM sürümünden gelmeyebileceğini gösterir. Scale-aware crop,
magnification, oriented bbox bilgisini kullanan prompt ve boundary-aware
refinement güçlü adaylardır. Ancak bu teknikler öğretmen-referans yanlılığını
çözmez; yine bağımsız referans gerekir.

### 5.1 Sınırlılıklar

- Ortak insan denetimi 128 görüntünün 126'sını ve 1.375 tile-instance'ın
  1.033'ünü kapsar. Örtüşen crop'lar nedeniyle bunlar 770 benzersiz insan
  anotasyonlu uçağa karşılık gelir.
- iSAID testi resmi insan etiketli validation havuzundan, SAMRS testi ise
  source-scene grouped ayrımdan gelir; bu nedenle veri setleri arası mutlak
  skor farkı yalnız referans kaynağına bağlanamaz.
- iSAID ve SAMRS annotation protokolleri insan/pseudo kaynağın yanında kendi
  tanım farklarına da sahip olabilir.
- GT prompt provenance'ı da farklıdır: iSAID resmi insan polygon envelope'u,
  SAMRS özgün DOTA detection RHBox'ıdır.
- YOLO hyperparameter'ları ve epok sayısı aynıdır; ancak eğitim görüntüsü
  sayıları farklı olduğu için 100 epok iki veri setinde aynı optimizer adımı
  sayısına karşılık gelmez. Epok-temelli warm-up'taki ilk üç kaydedilmiş
  öğrenme oranı da batch sayısına bağlı olarak çok küçük farklılaşır;
  dördüncü epoktan sonra çizelge aynıdır. Bu nedenle cross-dataset detector
  farkı yalnız veri zorluğuna bağlanamaz.
- Ana çalışma tek `plane` sınıfına odaklanır.
- Üç modelin sıralaması ortak audit subset'inde değişmemiştir.
- Sonuçlar pseudo-label ile eğitimin faydasını doğrudan ölçmez.

## 6. Sonuç

Teacher-generated maskeler eğitim verisi ve benchmark referansı olarak aynı
statüde değerlendirilmemelidir. Aynı görüntü ve tahmin üzerinde yalnız
referansın değiştirilmesi SAM1 skorunu yaklaşık 0,35 IoU yükseltmiştir.
Uzaktan algılama segmentation benchmark'ları üretici provenance'ı, bağımsız
insan denetimi ve eşleştirilmiş referans duyarlılığı raporlamalıdır.

## Kaynaklar

1. Wang et al. SAMRS: Scaling-up Remote Sensing Segmentation Dataset with
   Segment Anything Model. NeurIPS Datasets and Benchmarks, 2023.
2. Kirillov et al. Segment Anything. ICCV, 2023.
3. Zamir et al. iSAID: A Large-scale Dataset for Instance Segmentation in Aerial
   Images. CVPR Workshops, 2019.
4. Brachmann et al. On the Limits of Pseudo Ground Truth in Visual Camera
   Re-localisation. ICCV, 2021.
5. Arazo et al. Pseudo-Labeling and Confirmation Bias in Deep Semi-Supervised
   Learning. IJCNN, 2020.
6. Warfield et al. Validation of Image Segmentation by Estimating Rater Bias
   and Variance. Philosophical Transactions A, 2008.
7. SOPSeg: Prompt-based Small Object Instance Segmentation in Remote Sensing.
   arXiv:2509.03002, 2025.
"""


def rgb_hex(red: int, green: int, blue: int) -> str:
    return f"{red:02X}{green:02X}{blue:02X}"


def metric_fill(value: float) -> str:
    red, green, blue, _ = plt.get_cmap("RdYlGn")(Normalize(0.0, 1.0)(value))
    return rgb_hex(int(red * 255), int(green * 255), int(blue * 255))


def shade_cell(cell: Any, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def add_docx_table(document: Document, frame: pd.DataFrame) -> None:
    table = document.add_table(rows=1, cols=len(frame.columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for column_index, column in enumerate(frame.columns):
        cell = table.rows[0].cells[column_index]
        cell.text = str(column)
        shade_cell(cell, "D9EAF7")
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(8)
    for _, row in frame.iterrows():
        cells = table.add_row().cells
        for column_index, column in enumerate(frame.columns):
            value = row[column]
            if isinstance(value, (float, np.floating)):
                text = f"{float(value):.3f}"
                if 0.0 <= float(value) <= 1.0:
                    shade_cell(cells[column_index], metric_fill(float(value)))
            else:
                text = str(value)
            cells[column_index].text = text
            cells[column_index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cells[column_index].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(7.5)


def add_docx_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(item)


def build_docx(
    path: Path,
    evidence: dict[str, Any],
    *,
    issues: list[str],
    figures: dict[str, Path],
) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(9.5)

    title = document.add_heading(TITLE, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = document.add_paragraph(
        "Araştırma bildirisi taslağı | teacher_reference_bias_v1"
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if issues:
        warning = document.add_paragraph(PARTIAL_WARNING)
        warning.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in warning.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(180, 0, 0)

    shared = shared_reference_table(
        evidence["shared_summary"],
        evidence["shared_ci"],
    )
    strata_definitions = evidence["strata_definitions"]
    isaid_area_threshold = (
        100.0 * float(strata_definitions["isaid_plane"]["area_threshold"])
    )
    samrs_area_threshold = (
        100.0
        * float(strata_definitions["samrs_sota_plane"]["area_threshold"])
    )
    abstract = (
        "Otomatik pseudo-maskeler segmentation pretraining için yararlıdır; "
        "ancak bağımsız test ground truth'u olarak kullanıldıklarında referansı "
        "üreten modeli kayırabilir. iSAID ve SAMRS SOTA üzerinde eşlenmiş "
        "SAM1/SAM2/SAM3 deneyleri ile 1.033 ortak instance'ın dual-reference "
        f"analizi yapılmıştır. SAM1 IoU'su insan referansında "
        f"{shared.iloc[0]['İnsan IoU']:.3f}, kendi pseudo referansında "
        f"{shared.iloc[0]['SAM1 pseudo IoU']:.3f} bulunmuştur."
    )
    document.add_heading("Öz", level=1)
    document.add_paragraph(abstract)

    document.add_heading("1. Giriş ve Katkılar", level=1)
    document.add_paragraph(
        "Çalışma pseudo-label'ın eğitim faydası ile pseudo-reference'ın "
        "benchmark geçerliliğini ayırır. Ana soru, aynı tahmin üzerinde "
        "yalnız referans değiştiğinde skorun ne kadar değiştiğidir."
    )
    add_docx_bullets(
        document,
        [
            "Source-scene-safe ve iki veri setinde eşlenmiş plane protokolü.",
            "Aynı tahmin için insan ve SAM1 pseudo eşleştirilmiş değerlendirme.",
            "10.000 kaynak-sahne bootstrap ile referans enflasyonu aralıkları.",
            "SAM1, SAM2 ve SAM3 için ayrı GT-bbox ve YOLO-bbox analizi.",
        ],
    )

    document.add_heading("2. İlgili Çalışmalar", level=1)
    document.add_paragraph(
        "SAMRS detection annotation'larını SAM1 ViT-H ile segmentation "
        "maskesine dönüştürür ve pretraining verisi sağlar. Brachmann et al. "
        "pseudo-ground-truth üreten algoritmanın benzer yöntemleri "
        "kayırabildiğini göstermiştir. SOPSeg ise küçük uzaktan algılama "
        "nesnelerinde en büyük katkının crop/magnification ve oriented prompt "
        "tasarımından geldiğini raporlar."
    )

    document.add_heading("3. Deney Tasarımı", level=1)
    add_docx_bullets(
        document,
        [
            "İki veri setinde plane, 1024x1024 ve 128 test görüntüsü.",
            "Dört overlap x mask area katmanında 32'şer görüntü.",
            (
                "Strata görüntü düzeyinde: No Overlap için maksimum bbox-pair "
                "IoU=0, Overlap için IoU>=0,001."
            ),
            (
                "Low/high alan eşiği veri seti içi medyandır: iSAID insan "
                f"maskesinde %{isaid_area_threshold:.3f}, SAMRS pseudo "
                f"maskesinde %{samrs_area_threshold:.3f}; bu kırılım "
                "betimleyicidir."
            ),
            "Train/validation/test source scene kesişimi sıfır.",
            (
                "İnsan maskeleri kayıpsız RLE'dir; boş maske ve piksel alanı "
                "uyuşmazlığı sıfırdır."
            ),
            "Original detection bbox ve YOLO bbox ayrı koşullar.",
            (
                "iSAID GT prompt'u resmi insan bbox'ı, SAMRS GT prompt'u "
                "özgün DOTA RHBox'ıdır; pseudo-maskeden bbox üretilmez."
            ),
            "Ana değerlendirme instance-level; detector AP ayrı.",
        ],
    )

    document.add_heading("4. GT-bbox Sonuçları", level=1)
    add_docx_table(document, gt_bbox_table(evidence["aggregates"]))
    document.add_picture(str(figures["shared"]), width=Inches(6.55))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_docx_table(document, shared)

    document.add_heading("5. Detector ve YOLO-bbox Sonuçları", level=1)
    add_docx_table(document, detector_table(evidence["detector"]))
    add_docx_table(document, yolo_segmentation_table(evidence["segmentation"]))
    document.add_picture(str(figures["strata"]), width=Inches(6.55))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_heading("6. Tartışma ve Sonuç", level=1)
    document.add_paragraph(
        "SAM1'in kendi pseudo referansında yaklaşık kusursuz görünmesi, "
        "bağımsız insan sınırlarına göre kusursuz olduğu anlamına gelmez. "
        "Aynı 1.033 tahminde yalnız referansın değiştirilmesi SAM1 için "
        "yaklaşık 0,35 IoU artışı üretmiştir. SAMRS weak supervision için "
        "yararlı olabilir; benchmark kullanımı üretici provenance'ı ve insan "
        "denetimi gerektirir."
    )
    document.add_heading("Sınırlılıklar", level=2)
    add_docx_bullets(
        document,
        [
            "Ortak audit 1.375 instance'ın 1.033'ünü kapsar.",
            "Ana çalışma tek plane sınıfıyla sınırlıdır.",
            "Ortak subset'te model sıralaması değişmemiştir.",
            (
                "SAMRS low/high strata pseudo-mask alanından türediği için "
                "strata farkları referans etkisinin nedensel kanıtı değildir."
            ),
            (
                "GT bbox provenance'ı iSAID insan annotation'ı ve SAMRS "
                "DOTA detection RHBox'ı olarak farklıdır."
            ),
            (
                "YOLO ayarları ve epok sayısı aynıdır; eğitim görüntüsü "
                "sayıları farklı olduğundan toplam optimizer adımı aynı "
                "değildir. İlk üç warm-up epokundan sonra öğrenme oranı "
                "çizelgesi birebirdir."
            ),
            "Çalışma pseudo-label eğitim yararını doğrudan ölçmez.",
        ],
    )

    document.add_heading("Kaynaklar", level=1)
    references = [
        "Wang et al., SAMRS, NeurIPS Datasets and Benchmarks, 2023.",
        "Kirillov et al., Segment Anything, ICCV, 2023.",
        "Zamir et al., iSAID, CVPR Workshops, 2019.",
        "Brachmann et al., On the Limits of Pseudo Ground Truth, ICCV, 2021.",
        "Arazo et al., Pseudo-Labeling and Confirmation Bias, IJCNN, 2020.",
        "Warfield et al., Validation of Image Segmentation, 2008.",
        "SOPSeg, arXiv:2509.03002, 2025.",
    ]
    for reference in references:
        document.add_paragraph(reference)
    document.save(path)


def page_base(pdf: PdfPages, page_number: int, title: str) -> plt.Figure:
    figure = plt.figure(figsize=(8.27, 11.69), facecolor="white")
    figure.text(
        0.06,
        0.965,
        title,
        fontsize=12,
        fontweight="bold",
        color="#17324D",
        va="top",
    )
    figure.text(
        0.06,
        0.025,
        SHORT_TITLE,
        fontsize=6.8,
        color="#666666",
        ha="left",
    )
    figure.text(
        0.94,
        0.025,
        f"{page_number}/6",
        fontsize=6.8,
        color="#666666",
        ha="right",
    )
    figure.lines.append(
        plt.Line2D(
            [0.06, 0.94],
            [0.945, 0.945],
            transform=figure.transFigure,
            color="#A7B8C8",
            linewidth=0.7,
        )
    )
    return figure


def wrapped(text: str, width: int) -> str:
    paragraphs = []
    for paragraph in text.split("\n"):
        if not paragraph.strip():
            paragraphs.append("")
        else:
            paragraphs.append(
                textwrap.fill(
                    paragraph.strip(),
                    width=width,
                    break_long_words=False,
                    break_on_hyphens=False,
                )
            )
    return "\n".join(paragraphs)


def add_text(
    figure: plt.Figure,
    *,
    x: float,
    y: float,
    text: str,
    width: int,
    fontsize: float = 8.2,
    weight: str = "normal",
    color: str = "#222222",
) -> None:
    figure.text(
        x,
        y,
        wrapped(text, width),
        fontsize=fontsize,
        fontweight=weight,
        color=color,
        va="top",
        ha="left",
        linespacing=1.25,
    )


def add_table_axis(
    figure: plt.Figure,
    frame: pd.DataFrame,
    *,
    bbox: tuple[float, float, float, float],
    fontsize: float = 6.4,
) -> None:
    axis = figure.add_axes(bbox)
    axis.axis("off")
    if frame.empty:
        axis.text(
            0.5,
            0.5,
            "Sonuçlar henüz tamamlanmadı.",
            ha="center",
            va="center",
            color="#A00000",
            fontsize=8,
        )
        return
    display = frame.copy()
    for column in display.columns:
        if pd.api.types.is_float_dtype(display[column]):
            display[column] = display[column].map(lambda value: f"{value:.3f}")
    table = axis.table(
        cellText=display.values,
        colLabels=display.columns,
        cellLoc="center",
        loc="center",
        bbox=(0, 0, 1, 1),
    )
    table.auto_set_font_size(False)
    table.set_fontsize(fontsize)
    for (row, column), cell in table.get_celld().items():
        cell.set_linewidth(0.35)
        cell.set_edgecolor("#8A9AA8")
        if row == 0:
            cell.set_facecolor("#DCEAF4")
            cell.get_text().set_weight("bold")
        elif row % 2 == 0:
            cell.set_facecolor("#F4F7F9")


def add_image(
    figure: plt.Figure,
    image_path: Path,
    *,
    bbox: tuple[float, float, float, float],
) -> None:
    axis = figure.add_axes(bbox)
    axis.imshow(plt.imread(image_path))
    axis.axis("off")


def build_pdf(
    path: Path,
    evidence: dict[str, Any],
    *,
    issues: list[str],
    figures: dict[str, Path],
) -> None:
    gt = gt_bbox_table(evidence["aggregates"])
    shared = shared_reference_table(
        evidence["shared_summary"],
        evidence["shared_ci"],
    )
    detector = detector_table(evidence["detector"])
    yolo = yolo_segmentation_table(evidence["segmentation"])
    detector_pdf = detector.copy()
    if "Veri" in detector_pdf:
        detector_pdf["Veri"] = detector_pdf["Veri"].replace(
            {
                "iSAID Plane": "iSAID\nPlane",
                "SAMRS SOTA Plane": "SAMRS SOTA\nPlane",
            }
        )
    yolo_pdf = yolo.copy()
    if "Veri / referans" in yolo_pdf:
        yolo_pdf["Veri / referans"] = yolo_pdf["Veri / referans"].replace(
            {
                "iSAID Plane / İnsan": "iSAID / İnsan",
                "iSAID Plane / SAM1 pseudo": "iSAID / SAM1\npseudo",
                "SAMRS SOTA Plane / SAM1 pseudo": "SAMRS / SAM1\npseudo",
            }
        )
    quality = evidence["shared_quality"]
    quality_iou = float(
        quality[quality["stratum"] == "overall"].iloc[0]["mean_iou"]
    )
    unique_sensitivity = evidence["shared_unique_objects"].set_index("model")
    strata_definitions = evidence["strata_definitions"]
    isaid_area_threshold = (
        100.0 * float(strata_definitions["isaid_plane"]["area_threshold"])
    )
    samrs_area_threshold = (
        100.0
        * float(strata_definitions["samrs_sota_plane"]["area_threshold"])
    )

    with PdfPages(path) as pdf:
        page = page_base(pdf, 1, "Araştırma Bildirisi")
        add_text(
            page,
            x=0.08,
            y=0.915,
            width=61,
            fontsize=13.0,
            weight="bold",
            color="#17324D",
            text=TITLE,
        )
        page.text(
            0.08,
            0.835,
            "teacher_reference_bias_v1 | " + date.today().isoformat(),
            ha="left",
            fontsize=8,
            color="#536979",
        )
        if issues:
            page.text(
                0.08,
                0.812,
                PARTIAL_WARNING,
                ha="left",
                fontsize=7,
                color="#A00000",
                fontweight="bold",
            )
        add_text(
            page,
            x=0.08,
            y=0.775,
            width=104,
            fontsize=8.5,
            text=(
                "Öz. Otomatik pseudo-maskeler segmentation pretraining için "
                "değerlidir; ancak bağımsız test ground truth'u olarak "
                "kullanıldıklarında referansı üreten modeli kayırabilir. "
                "iSAID ve SAMRS SOTA üzerinde aynı plane sınıfı, aynı 128 "
                "test görüntüsü ve aynı SAM1/SAM2/SAM3 bbox-prompted pipeline "
                "ile kontrollü deney yapılmıştır. 126 SAMRS tile'ı iSAID "
                "insan etiketli kaynaklarına piksel düzeyinde eşlenmiş ve "
                "1.033 ortak instance için aynı tahminler iki referansa "
                f"karşı ölçülmüştür. SAM1 IoU'su insan referansında "
                f"{shared.iloc[0]['İnsan IoU']:.3f}, kendi pseudo referansında "
                f"{shared.iloc[0]['SAM1 pseudo IoU']:.3f} bulunmuştur. "
                "Bulgular pseudo-label eğitim yararı ile benchmark "
                "validity kavramlarının ayrılması gerektiğini gösterir."
            ),
        )
        add_image(page, figures["shared"], bbox=(0.08, 0.31, 0.84, 0.34))
        add_text(
            page,
            x=0.08,
            y=0.265,
            width=104,
            fontsize=8.1,
            weight="bold",
            color="#17324D",
            text=(
                "Katkılar: kaynak-sahne güvenli eşlenmiş protokol; aynı "
                "tahminle çift referanslı değerlendirme; 10.000 kümeli "
                "bootstrap; SAM1/SAM2/SAM3 GT-bbox ve YOLO-bbox ayrımı."
            ),
        )
        add_text(
            page,
            x=0.08,
            y=0.175,
            width=104,
            fontsize=7.6,
            text=(
                "Anahtar kelimeler: remote sensing, Segment Anything, "
                "pseudo-label, evaluation bias, instance segmentation, "
                "öğretmen-referans yakınlığı."
            ),
        )
        pdf.savefig(page)
        plt.close(page)

        page = page_base(pdf, 2, "1. Giriş ve 2. İlgili Çalışmalar")
        add_text(
            page,
            x=0.06,
            y=0.92,
            width=52,
            fontsize=8.2,
            weight="bold",
            color="#17324D",
            text="1. Giriş",
        )
        add_text(
            page,
            x=0.06,
            y=0.885,
            width=52,
            text=(
                "Pixel-level uzaktan algılama annotation'ı pahalıdır. SAMRS, "
                "mevcut detection bbox'larını SAM1 prompt'u olarak kullanıp "
                "büyük ölçekli maskeler üretir. Bu yaklaşım pretraining için "
                "ölçek sağlar. Ancak aynı öğretmen kendi ürettiği maskelere "
                "karşı ölçüldüğünde, nesne sınırına uyum ile öğretmenin karar "
                "stilini yeniden üretme etkisi karışır.\n\n"
                "İlk iSAID ve SAMRS deneylerimiz arasında büyük fark görülmesi "
                "bu soruyu doğurdu. Basit veri seti karşılaştırması yeterli "
                "değildir; görüntü zorluğu, annotation protokolü ve model "
                "referansı birlikte değişir. Bu nedenle aynı tahmini iki "
                "referansta ölçen kontrollü tasarım gerekir.\n\n"
                "Çalışmanın katkıları: (i) iki veri setinde eşlenmiş plane "
                "protokolü, (ii) insan/pseudo eşleştirilmiş değerlendirme, "
                "(iii) kaynak-sahne bootstrap, (iv) detector ve segmenter "
                "etkisinin ayrılması, (v) üretici provenance denetimi."
            ),
        )
        add_text(
            page,
            x=0.53,
            y=0.92,
            width=52,
            fontsize=8.2,
            weight="bold",
            color="#17324D",
            text="2. İlgili Çalışmalar",
        )
        add_text(
            page,
            x=0.53,
            y=0.885,
            width=52,
            text=(
                "SAMRS, DOTA, DIOR, FAIR1M ve HRSC2016 detection "
                "annotation'larını SAM ile maskeye çevirir. Resmi SOTA-RBB "
                "kodu SAM1 ViT-H ve rotated kutunun minimum horizontal "
                "çevreleyicisi RHBox kullanır. Makalenin ana downstream "
                "iddiası pretraining sonrası bağımsız iSAID/Potsdam "
                "fine-tuning sonuçlarıdır.\n\n"
                "Brachmann et al. pseudo-ground-truth üreten SfM/SLAM "
                "algoritmalarının benzer localisation yöntemlerini "
                "kayırabildiğini gösterir. Bu çalışma segmentation alanındaki "
                "en yakın metodolojik öncüldür. Arazo et al. confirmation "
                "bias'ı eğitim dinamiği olarak inceler; bizim problemimiz "
                "ölçüm geçerliliği problemidir.\n\n"
                "SOPSeg ablation'ında region-adaptive magnification yaklaşık "
                "+7,75 mIoU, oriented prompt +2,29 ve edge-aware decoder +1,27 "
                "kazandırır. Bu sonuç scale-aware preprocessing ve prompt "
                "geometrisinin yeni model sürümü kadar önemli olabileceğini "
                "gösterir; bağımsız ground truth gereksinimini ortadan "
                "kaldırmaz."
            ),
        )
        add_text(
            page,
            x=0.06,
            y=0.52,
            width=52,
            fontsize=8.2,
            weight="bold",
            color="#17324D",
            text="2.1 SAMRS'nin amaçlanan kullanımı",
        )
        add_text(
            page,
            x=0.06,
            y=0.485,
            width=52,
            text=(
                "SAMRS makalesi pseudo-maskeleri bağımsız bir insan benchmark'ı "
                "olarak değil, büyük ölçekli segmentation pretraining kaynağı "
                "olarak konumlandırır. Bu kullanımda öğretmen etiketiyle öğrenilen "
                "temsilin gerçek faydası, sonradan iSAID veya Potsdam gibi "
                "bağımsız insan etiketli downstream testlerde ölçülür. Bizim "
                "eleştirimiz veri üretimine değil, aynı pseudo-maskenin test "
                "ground truth'u gibi yorumlanmasına yöneliktir."
            ),
        )
        add_text(
            page,
            x=0.53,
            y=0.52,
            width=52,
            fontsize=8.2,
            weight="bold",
            color="#17324D",
            text="2.2 Eğitim yanlılığı ve ölçüm yanlılığı",
        )
        add_text(
            page,
            x=0.53,
            y=0.485,
            width=52,
            text=(
                "Confirmation bias, modelin kendi hatalı pseudo-label'larını "
                "eğitim sırasında pekiştirmesidir. Teacher-reference affinity "
                "ise eğitim yapılmasa bile ortaya çıkabilir: tahmin, kendisini "
                "üreten modelin sınır stiline benzeyen referansla ölçülünce skor "
                "yükselir. Bu bildiride tahmin sabit tutulduğu için ölçülen "
                "etki doğrudan değerlendirme referansından gelir."
            ),
        )
        add_table_axis(
            page,
            pd.DataFrame(
                [
                    ["SAMRS", "Pseudo-mask üretimi", "Ön eğitim"],
                    ["Brachmann et al.", "Pseudo-GT yöntem yakınlığı", "Değerlendirme"],
                    ["Arazo et al.", "Pseudo-label confirmation bias", "Eğitim"],
                    ["Bu çalışma", "Aynı tahmin, iki referans", "Değerlendirme"],
                ],
                columns=["Çalışma", "Ana sorun", "Bağlam"],
            ),
            bbox=(0.10, 0.12, 0.80, 0.17),
            fontsize=6.8,
        )
        pdf.savefig(page)
        plt.close(page)

        page = page_base(pdf, 3, "3. Materyal ve Yöntem")
        add_text(
            page,
            x=0.06,
            y=0.92,
            width=52,
            weight="bold",
            color="#17324D",
            text="3.1 Veri ve provenance",
        )
        add_text(
            page,
            x=0.06,
            y=0.885,
            width=52,
            text=(
                "iSAID plane maskeleri insan tarafından bağımsız çizilmiş "
                "polygon'lardır; tile maskeleri kayıpsız COCO RLE'dir ve "
                "boş maske/alan denetimi sıfırdır. SAMRS SOTA plane maskeleri "
                "SAM1 ViT-H ve "
                "original DOTA RHBox prompt'larıyla üretilmiştir. Resmi "
                "17.555 dosya ve 615.407 instance için numeric class ID, RBox "
                "ve RHBox karşılaştırılmış; geometri farkı bulunmamıştır.\n\n"
                "İki veri setinde giriş 1024x1024, hedef plane ve test sayısı "
                "128'dir. Testte no-overlap/overlap ile low/high mask area "
                "çarpımından dört strata vardır ve her biri 32 görüntüdür. "
                "Strata görüntü düzeyindedir: no-overlap için maksimum bbox "
                "çifti IoU=0, overlap için IoU>=0,001'dir. Low/high eşiği "
                f"veri seti içi medyandır; iSAID insan maskesinde "
                f"%{isaid_area_threshold:.3f}, SAMRS pseudo maskesinde "
                f"%{samrs_area_threshold:.3f}'tür. Bu kırılım betimleyicidir. "
                "Split birimi tile değil parent source scene'dir; train, "
                "validation ve test kesişimi sıfırdır."
            ),
        )
        add_text(
            page,
            x=0.53,
            y=0.92,
            width=52,
            weight="bold",
            color="#17324D",
            text="3.2 Pipeline ve değerlendirme",
        )
        add_text(
            page,
            x=0.53,
            y=0.885,
            width=52,
            text=(
                "SAM1 ViT-H, SAM2.1 Hiera Large ve SAM3 aynı görüntü ve aynı "
                "bbox ile çalışır. iSAID GT prompt'u resmi insan bbox'ı, "
                "SAMRS GT prompt'u özgün DOTA RHBox'ıdır; SAM1 pseudo-maskeden "
                "bbox üretilmez. Üç seed ile eğitilmiş YOLO bbox ayrı "
                "koşuldur.\n\n"
                "Ana metrik instance-level mask IoU'dur. Dice, pixel "
                "precision, pixel recall, Boundary IoU ve "
                "Success@0.50/0.75/0.90 destekleyicidir. Detector için COCO "
                "bbox AP50/AP75/AP90/AP50-95 ayrı hesaplanır. Her seed'in "
                "YOLO confidence eşiği validation setinde bbox IoU 0,50 için "
                "en yüksek F1 noktasında dondurulur; test eşik seçimine "
                "girmez. YOLO koşulunda "
                "eşleşmeyen GT boş maskeyle sıfır skor alır; eşleşmeyen "
                "detector tahmini detector AP ve image-union hesabında "
                "korunur. Instance tablosu COCO mask AP değildir.\n\n"
                "Güven aralığı kaynak-sahne kümeli 10.000 bootstrap ile "
                "hesaplanır. Model farkları kaynak-sahne ortalamaları üzerinde "
                "eşleştirilmiş Wilcoxon testi ve Holm düzeltmesiyle incelenir."
            ),
        )
        add_text(
            page,
            x=0.06,
            y=0.49,
            width=104,
            weight="bold",
            color="#17324D",
            text="3.3 Kontrollü ve ortak görüntü çift-referans tasarımı",
        )
        add_text(
            page,
            x=0.06,
            y=0.455,
            width=104,
            text=(
                "Kontrollü iSAID deneyinde SAM1 GT-bbox tahmini ikinci "
                "pseudo referans olarak dondurulur. Ortak görüntü audit'inde "
                "SAMRS tile'ları iSAID kaynaklarına template matching ile geri "
                "eşlenir; yalnız template score >= 0,995 ve piksel eşitliği "
                "sağlanan tile'lar kabul edilir. SAMRS ve iSAID plane "
                "instance'ları one-to-one bbox IoU >= 0,50 ile eşlenir. "
                "Tahmin değişmez; insan ve SAM1 pseudo referans sırayla "
                "kullanılır. Bu tasarım veri seti zorluğunu sabitler. Eşleşen "
                "1.033 tile-instance görünümü, örtüşen crop'lar nedeniyle 770 "
                "benzersiz insan-anotasyonlu uçağa karşılık gelir."
            ),
        )
        add_table_axis(
            page,
            pd.DataFrame(
                [
                    ["iSAID", 128, 1045, 45, "İnsan"],
                    ["SAMRS SOTA", 128, 1375, 42, "SAM1 pseudo"],
                    ["Ortak denetim", 126, "1033 / 770", 35, "İnsan + pseudo"],
                ],
                columns=["Koşul", "Görüntü", "Görünüm / nesne", "Sahne", "Referans"],
            ),
            bbox=(0.10, 0.15, 0.80, 0.20),
            fontsize=7.2,
        )
        pdf.savefig(page)
        plt.close(page)

        page = page_base(pdf, 4, "4. Sonuçlar: GT-bbox ve Referans Etkisi")
        add_text(
            page,
            x=0.06,
            y=0.92,
            width=104,
            text=(
                f"Official SAMRS pseudo-mask ile bağımsız iSAID insan maskesi "
                f"arasındaki ortalama IoU {quality_iou:.3f}'tür. Referans "
                "yüksek kaliteli, fakat insan ground truth ile özdeş değildir. "
                "Aynı 1.033 tile-instance görünümü 770 benzersiz uçağa aittir. "
                "Üç modelin de pseudo skoru yükselmiş, artış öğretmen SAM1 "
                "için en büyük olmuştur."
            ),
        )
        add_table_axis(page, shared, bbox=(0.08, 0.69, 0.84, 0.16), fontsize=7.0)
        add_image(page, figures["shared"], bbox=(0.08, 0.33, 0.84, 0.32))
        add_text(
            page,
            x=0.06,
            y=0.28,
            width=104,
            text=(
                "İnsan ve pseudo referans model sırası ortak alt kümede "
                f"değişmemiştir (Spearman = "
                f"{float(evidence['shared_ranking']['spearman_correlation']):.1f}; "
                f"Kendall tau = "
                f"{float(evidence['shared_ranking']['kendall_tau']):.1f}). "
                "Bu nedenle ana iddia ranking reversal değil, model ailesine "
                "göre farklılaşan skor enflasyonudur. Kontrollü iSAID "
                "referans değişimi de aynı yönlü sonucu verir. Benzersiz uçak "
                "başına duyarlılık analizinde SAM1/SAM2/SAM3 IoU enflasyonu "
                f"{float(unique_sensitivity.loc['sam1', 'mean_iou_inflation']):.3f}/"
                f"{float(unique_sensitivity.loc['sam2', 'mean_iou_inflation']):.3f}/"
                f"{float(unique_sensitivity.loc['sam3', 'mean_iou_inflation']):.3f}'tür."
            ),
        )
        pdf.savefig(page)
        plt.close(page)

        page = page_base(pdf, 5, "4. Detector ve YOLO-bbox Sonuçları")
        add_text(
            page,
            x=0.06,
            y=0.92,
            width=104,
            weight="bold",
            color="#17324D",
            text="4.3 Üç seed YOLO detector sonuçları",
        )
        add_table_axis(
            page,
            detector_pdf,
            bbox=(0.08, 0.74, 0.84, 0.14),
            fontsize=7.0,
        )
        add_text(
            page,
            x=0.06,
            y=0.70,
            width=104,
            weight="bold",
            color="#17324D",
            text="4.4 YOLO-bbox segmentation",
        )
        add_table_axis(
            page,
            yolo_pdf,
            bbox=(0.06, 0.42, 0.88, 0.25),
            fontsize=5.7,
        )
        add_image(page, figures["strata"], bbox=(0.07, 0.08, 0.86, 0.29))
        pdf.savefig(page)
        plt.close(page)

        page = page_base(pdf, 6, "5. Tartışma, Sınırlılıklar ve Sonuç")
        add_text(
            page,
            x=0.06,
            y=0.92,
            width=52,
            weight="bold",
            color="#17324D",
            text="5. Tartışma",
        )
        add_text(
            page,
            x=0.06,
            y=0.885,
            width=52,
            text=(
                "SAM1'in kendi pseudo referansında yaklaşık kusursuz görünmesi "
                "bağımsız insan sınırlarına göre kusursuz olduğu anlamına "
                "gelmez. Aynı tahminde yalnız referans değişimi SAM1 "
                "için yaklaşık 0,35 IoU artışı üretmiştir. SAM2 ve SAM3 de "
                "yükselmiş, fakat artış öğretmen SAM1 için daha büyüktür.\n\n"
                "Sonuç SAMRS'nin değersiz olduğunu göstermez. Pseudo-maskeler "
                "pretraining, distillation ve weak supervision için yararlı "
                "olabilir. Downstream fayda bağımsız insan etiketli test "
                "setinde ölçülmelidir. Benchmark kullanımı üretici model, "
                "checkpoint, prompt ve insan denetimi provenance'ı gerektirir.\n\n"
                "Remote sensing küçük nesnelerinde SOPSeg bulguları crop, "
                "magnification, oriented prompt ve boundary refinement "
                "tekniklerinin önemli olduğunu gösterir. Bu teknikler yeni "
                "deney yönüdür, ancak bağımsız referans gereksinimini çözmez."
            ),
        )
        add_text(
            page,
            x=0.53,
            y=0.92,
            width=52,
            weight="bold",
            color="#17324D",
            text="Sınırlılıklar ve Sonuç",
        )
        add_text(
            page,
            x=0.53,
            y=0.885,
            width=52,
            text=(
                "Ortak audit 1.375 tile-instance'ın 1.033'ünü kapsar; bunlar "
                "770 benzersiz insan-anotasyonlu uçağa karşılık gelir. Ana deney "
                "tek plane sınıfındadır. iSAID testi resmi validation "
                "havuzundan, SAMRS testi grouped source-scene ayrımından "
                "gelir; cross-dataset mutlak fark yalnız referans etkisi "
                "değildir. Annotation protokollerinin kaynak "
                "türünden bağımsız tanım ve GT bbox provenance farkları "
                "olabilir. Ortak subset'te "
                "sıralama değişmemiştir. YOLO ayarları ve epok sayısı aynı "
                "olsa da train boyutları farklı olduğundan toplam optimizer "
                "adımı aynı değildir; ilk üç warm-up epokundan sonra LR "
                "çizelgesi birebirdir. "
                "Çalışma pseudo-label training "
                "utility'yi doğrudan ölçmez.\n\n"
                "Sonuç olarak öğretmen üretimli maskeler eğitim verisi ve "
                "benchmark referansı olarak aynı statüde değerlendirilmemelidir. "
                "Aynı tahmin üzerinde SAM1 için yaklaşık 0,35 IoU "
                "enflasyonu ölçülmüştür. Uzaktan algılama segmentation "
                "benchmark'ları üretici provenance'ı, bağımsız insan denetimi "
                "ve eşleştirilmiş referans duyarlılığı raporlamalıdır.\n\n"
                "Kaynaklar\n"
                "[1] Wang et al., SAMRS, NeurIPS D&B 2023.\n"
                "[2] Kirillov et al., Segment Anything, ICCV 2023.\n"
                "[3] Zamir et al., iSAID, CVPRW 2019.\n"
                "[4] Brachmann et al., Limits of Pseudo GT, ICCV 2021.\n"
                "[5] Arazo et al., Confirmation Bias, IJCNN 2020.\n"
                "[6] Warfield et al., Segmentation Validation, 2008.\n"
                "[7] SOPSeg, arXiv:2509.03002, 2025."
            ),
        )
        add_text(
            page,
            x=0.06,
            y=0.50,
            width=104,
            weight="bold",
            color="#17324D",
            text="Benchmark için önerilen asgari raporlama",
        )
        add_table_axis(
            page,
            pd.DataFrame(
                [
                    ["Referans kaynağı", "İnsan / pseudo ve üretici açıkça yazılmalı"],
                    ["Üretici provenance", "Model, checkpoint, prompt ve eşikler verilmeli"],
                    ["Bağımsız denetim", "Temsili subset insan maskesiyle ölçülmeli"],
                    ["Duyarlılık analizi", "Aynı tahmin iki referansta raporlanmalı"],
                    ["Split güvenliği", "Kaynak sahne kesişimi sıfır olmalı"],
                ],
                columns=["Kontrol", "Gerekli raporlama"],
            ),
            bbox=(0.08, 0.245, 0.84, 0.20),
            fontsize=6.8,
        )
        add_text(
            page,
            x=0.06,
            y=0.19,
            width=104,
            fontsize=8.4,
            weight="bold",
            color="#17324D",
            text=(
                "Temel sonuç: öğretmen üretimli maskeler eğitim verisi olarak "
                "yararlı olabilir; fakat bağımsız insan referansıyla aynı "
                "epistemik statüde bir test ground truth'u değildir."
            ),
        )
        pdf.savefig(page)
        plt.close(page)


def copy_figures(study_root: Path, output_dir: Path) -> dict[str, Path]:
    figure_dir = output_dir / "figures"
    figure_dir.mkdir(parents=True, exist_ok=True)
    sources = {
        "shared": study_root / "figures" / "shared_human_reference_comparison.png",
        "strata": study_root / "figures" / "gt_bbox_strata_heatmap.png",
        "reference": study_root / "figures" / "gt_bbox_reference_comparison.png",
    }
    outputs: dict[str, Path] = {}
    for key, source in sources.items():
        if not source.exists():
            raise FileNotFoundError(source)
        destination = figure_dir / source.name
        shutil.copy2(source, destination)
        outputs[key] = destination
    return outputs


def write_manifest(
    path: Path,
    *,
    inputs: list[Path],
    outputs: list[Path],
    issues: list[str],
) -> None:
    payload = {
        "schema_version": 1,
        "status": "completed" if not issues else "partial_draft",
        "completion_issues": issues,
        "inputs": [
            {"path": str(item), "sha256": sha256_file(item)}
            for item in sorted(set(inputs))
        ],
        "outputs": [
            {"path": str(item), "sha256": sha256_file(item)}
            for item in sorted(set(outputs))
        ],
    }
    path.write_text(
        json.dumps(payload, indent=2, ensure_ascii=False) + "\n",
        encoding="utf-8",
    )


def main() -> None:
    args = parse_args()
    evidence = load_evidence(args.study_root)
    issues = completion_issues(evidence)
    if issues and not args.allow_partial:
        raise RuntimeError(
            "Paper evidence is incomplete:\n- " + "\n- ".join(issues)
        )

    args.output_dir.mkdir(parents=True, exist_ok=True)
    figures = copy_figures(args.study_root, args.output_dir)
    figure_names = {
        key: str(path.relative_to(args.output_dir))
        for key, path in figures.items()
    }
    markdown_path = args.output_dir / "teacher_reference_bias_paper.md"
    docx_path = args.output_dir / "teacher_reference_bias_paper.docx"
    pdf_path = args.output_dir / "teacher_reference_bias_paper_6pages.pdf"
    markdown_path.write_text(
        build_markdown(
            evidence,
            issues=issues,
            figure_names=figure_names,
        ),
        encoding="utf-8",
    )
    build_docx(
        docx_path,
        evidence,
        issues=issues,
        figures=figures,
    )
    build_pdf(
        pdf_path,
        evidence,
        issues=issues,
        figures=figures,
    )
    outputs = [markdown_path, docx_path, pdf_path, *figures.values()]
    manifest_path = args.output_dir / "paper_manifest.json"
    write_manifest(
        manifest_path,
        inputs=evidence["input_paths"],
        outputs=outputs,
        issues=issues,
    )
    for output in [markdown_path, docx_path, pdf_path, manifest_path]:
        print(output)


if __name__ == "__main__":
    main()
