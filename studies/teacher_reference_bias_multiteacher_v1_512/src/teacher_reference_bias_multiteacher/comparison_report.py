from __future__ import annotations

import math
import re
import textwrap
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from matplotlib.backends.backend_pdf import PdfPages
from PIL import Image

from .paths import DATASETS, REFERENCE_TYPES, STRATA


REFERENCE_LABELS = {
    "human": "Human",
    "pseudo_sam1": "SAM1 pseudo",
    "pseudo_sam2": "SAM2 pseudo",
    "pseudo_sam3": "SAM3 pseudo",
}
DATASET_LABELS = {
    "isaid_plane": "iSAID Plane",
    "isaid_small_vehicle": "iSAID Small Vehicle",
}


def interpolate_color(value: float) -> str:
    value = max(0.0, min(1.0, value))
    red, yellow, green = (248, 105, 107), (255, 235, 132), (99, 190, 123)
    if value <= 0.5:
        ratio, left, right = value / 0.5, red, yellow
    else:
        ratio, left, right = (value - 0.5) / 0.5, yellow, green
    rgb = tuple(
        round(left[index] + (right[index] - left[index]) * ratio)
        for index in range(3)
    )
    return "".join(f"{channel:02X}" for channel in rgb)


def numeric(value: object) -> float | None:
    match = re.match(r"^\s*([0-9]+(?:\.[0-9]+)?)", str(value))
    if match is None:
        return None
    result = float(match.group(1))
    return result if 0 <= result <= 1 else None


def matrix_table(
    aggregates: pd.DataFrame,
    *,
    dataset_id: str,
    bbox_source: str,
) -> pd.DataFrame:
    selected = aggregates[
        (aggregates["dataset_id"] == dataset_id)
        & (aggregates["bbox_source"] == bbox_source)
        & (aggregates["stratum"] == "overall")
    ]
    pivot = selected.pivot_table(
        index="model",
        columns="reference_type",
        values="mean_iou",
        aggfunc="mean",
    ).loc[["sam1", "sam2", "sam3"], list(REFERENCE_TYPES)]
    output = pd.DataFrame({"Model": ["SAM1", "SAM2", "SAM3"]})
    for reference_type in REFERENCE_TYPES:
        output[REFERENCE_LABELS[reference_type]] = [
            f"{value:.3f}" for value in pivot[reference_type]
        ]
    return output


def main_metric_table(aggregates: pd.DataFrame) -> pd.DataFrame:
    overall = aggregates[aggregates["stratum"] == "overall"]
    rows = []
    for dataset_id in DATASETS:
        for bbox_source in ("gt_bbox", "yolo_bbox"):
            for reference_type in REFERENCE_TYPES:
                selected = overall[
                    (overall["dataset_id"] == dataset_id)
                    & (overall["bbox_source"] == bbox_source)
                    & (overall["reference_type"] == reference_type)
                ].sort_values(["mean_iou", "model"], ascending=[False, True])
                best = selected.iloc[0]
                rows.append(
                    {
                        "Dataset": DATASET_LABELS[dataset_id],
                        "BBox": "GT" if bbox_source == "gt_bbox" else "YOLO",
                        "Reference": REFERENCE_LABELS[reference_type],
                        "Best model": str(best["model"]).upper(),
                        "Avg IoU": f"{float(best['mean_iou']):.3f}",
                        "Avg Dice": f"{float(best['mean_dice']):.3f}",
                        "Avg Precision": f"{float(best['mean_precision']):.3f}",
                        "Avg Recall": f"{float(best['mean_recall']):.3f}",
                    }
                )
    return pd.DataFrame(rows)


def empty_reference_table() -> pd.DataFrame:
    rows = []
    for dataset_id, source in DATASETS.items():
        for teacher in ("sam1", "sam2", "sam3"):
            path = (
                source.predictions_root
                / teacher
                / "gt_bbox"
                / "predictions.jsonl"
            )
            statuses = pd.read_json(path, lines=True, typ="frame", dtype=False)[
                "status"
            ]
            empty = int((statuses == "empty_mask").sum())
            rows.append(
                {
                    "Dataset": DATASET_LABELS[dataset_id],
                    "Teacher": teacher.upper(),
                    "Instances": len(statuses),
                    "Empty masks": empty,
                    "Empty rate": f"{empty / len(statuses):.3f}",
                }
            )
    return pd.DataFrame(rows)


def strata_advantage_table(aggregates: pd.DataFrame) -> pd.DataFrame:
    rows = []
    for dataset_id in DATASETS:
        for stratum in STRATA:
            for teacher in ("sam1", "sam2", "sam3"):
                selected = aggregates[
                    (aggregates["dataset_id"] == dataset_id)
                    & (aggregates["bbox_source"] == "yolo_bbox")
                    & (aggregates["reference_type"] == f"pseudo_{teacher}")
                    & (aggregates["stratum"] == stratum)
                ].set_index("model")
                alternatives = [model for model in selected.index if model != teacher]
                advantage = float(
                    selected.loc[teacher, "mean_iou"]
                    - selected.loc[alternatives, "mean_iou"].mean()
                )
                rows.append(
                    {
                        "Dataset": DATASET_LABELS[dataset_id],
                        "Stratum": stratum.replace("__", " × ").replace("_", " ").title(),
                        "Teacher": teacher.upper(),
                        "YOLO teacher advantage": f"{advantage:+.3f}",
                    }
                )
    return pd.DataFrame(rows)


def agreement_table(agreement: pd.DataFrame) -> pd.DataFrame:
    output = agreement.copy()
    output["Dataset"] = output["dataset_id"].map(DATASET_LABELS)
    output["Reference A"] = output["reference_a"].map(REFERENCE_LABELS)
    output["Reference B"] = output["reference_b"].map(REFERENCE_LABELS)
    output["Mean instance IoU"] = output["mean_instance_iou"].map(
        lambda value: f"{value:.3f}"
    )
    output["Instances"] = output["instance_count"].astype(int)
    return output[
        ["Dataset", "Reference A", "Reference B", "Mean instance IoU", "Instances"]
    ]


def _markdown_table(frame: pd.DataFrame) -> str:
    headers = [str(column) for column in frame.columns]
    lines = [
        "| " + " | ".join(headers) + " |",
        "| " + " | ".join("---" for _ in headers) + " |",
    ]
    for _, row in frame.iterrows():
        lines.append("| " + " | ".join(str(row[column]) for column in frame.columns) + " |")
    return "\n".join(lines)


def _shade(cell, color: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    properties.append(shading)


def _docx_table(document: Document, frame: pd.DataFrame, title: str) -> None:
    document.add_heading(title, level=2)
    table = document.add_table(rows=1, cols=len(frame.columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, column in enumerate(frame.columns):
        cell = table.rows[0].cells[index]
        cell.text = str(column)
        _shade(cell, "1F4E79")
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True
            run.font.size = Pt(7)
    for _, row in frame.iterrows():
        cells = table.add_row().cells
        for index, column in enumerate(frame.columns):
            cells[index].text = str(row[column])
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cells[index].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = "DejaVu Sans"
                    run.font.size = Pt(6.7)
            if column in {
                "Human",
                "SAM1 pseudo",
                "SAM2 pseudo",
                "SAM3 pseudo",
                "Avg IoU",
                "Avg Dice",
                "Avg Precision",
                "Avg Recall",
                "Mean instance IoU",
                "Empty rate",
            }:
                value = numeric(row[column])
                if value is not None:
                    _shade(cells[index], interpolate_color(value))


def _docx_bullets(document: Document, bullets: list[str]) -> None:
    for bullet in bullets:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(bullet)
        run.font.name = "DejaVu Sans"
        run.font.size = Pt(9)


def _configure_docx(document: Document) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = section.bottom_margin = Inches(0.45)
    section.left_margin = section.right_margin = Inches(0.50)
    for style_name in ("Normal", "Title", "Heading 1", "Heading 2"):
        document.styles[style_name].font.name = "DejaVu Sans"


def _pdf_text_page(pdf: PdfPages, title: str, bullets: list[str]) -> None:
    figure = plt.figure(figsize=(11.69, 8.27))
    figure.text(0.06, 0.93, title, fontsize=18, fontweight="bold", va="top")
    y = 0.84
    for bullet in bullets:
        wrapped = textwrap.wrap(bullet, width=105)
        figure.text(0.075, y, "• " + wrapped[0], fontsize=10.5, va="top")
        y -= 0.040
        for line in wrapped[1:]:
            figure.text(0.093, y, line, fontsize=10.5, va="top")
            y -= 0.036
        y -= 0.020
    plt.axis("off")
    pdf.savefig(figure, bbox_inches="tight")
    plt.close(figure)


def _pdf_image_page(pdf: PdfPages, title: str, image_path: Path, note: str) -> None:
    figure = plt.figure(figsize=(11.69, 8.27))
    figure.text(0.05, 0.95, title, fontsize=16, fontweight="bold", va="top")
    figure.text(0.05, 0.89, note, fontsize=9.2, va="top")
    axis = figure.add_axes((0.04, 0.05, 0.92, 0.80))
    axis.imshow(Image.open(image_path))
    axis.axis("off")
    pdf.savefig(figure)
    plt.close(figure)


def _pdf_table_page(pdf: PdfPages, title: str, tables: list[tuple[str, pd.DataFrame]]) -> None:
    figure = plt.figure(figsize=(11.69, 8.27))
    figure.text(0.04, 0.96, title, fontsize=16, fontweight="bold", va="top")
    table_height = 0.78 / len(tables)
    for index, (subtitle, frame) in enumerate(tables):
        top = 0.87 - index * table_height
        figure.text(0.05, top + 0.025, subtitle, fontsize=10, fontweight="bold")
        axis = figure.add_axes((0.04, top - table_height + 0.04, 0.92, table_height - 0.05))
        axis.axis("off")
        table = axis.table(
            cellText=frame.astype(str).values,
            colLabels=list(frame.columns),
            cellLoc="center",
            loc="center",
        )
        table.auto_set_font_size(False)
        table.set_fontsize(7.2 if len(frame.columns) <= 6 else 6.4)
        table.scale(1, 1.35)
        for (row, column), cell in table.get_celld().items():
            if row == 0:
                cell.set_facecolor("#1F4E79")
                cell.get_text().set_color("white")
                cell.get_text().set_fontweight("bold")
            elif column > 0:
                value = numeric(frame.iloc[row - 1, column])
                if value is not None:
                    cell.set_facecolor("#" + interpolate_color(value))
    pdf.savefig(figure)
    plt.close(figure)


def write_comparison_report(
    *,
    output_dir: Path,
    aggregates: pd.DataFrame,
    agreement: pd.DataFrame,
    rankings: pd.DataFrame,
    matrix_figure: Path,
    effect_figure: Path,
) -> dict[str, Path]:
    output_dir.mkdir(parents=True, exist_ok=True)
    matrices = {
        (dataset_id, bbox): matrix_table(
            aggregates, dataset_id=dataset_id, bbox_source=bbox
        )
        for dataset_id in DATASETS
        for bbox in ("gt_bbox", "yolo_bbox")
    }
    metrics = main_metric_table(aggregates)
    empties = empty_reference_table()
    agreements = agreement_table(agreement)
    strata = strata_advantage_table(aggregates)
    rankings_display = rankings.copy()
    rankings_display["dataset_id"] = rankings_display["dataset_id"].map(DATASET_LABELS)
    rankings_display["reference_type"] = rankings_display["reference_type"].map(REFERENCE_LABELS)
    rankings_display = rankings_display.rename(
        columns={
            "dataset_id": "Dataset",
            "bbox_source": "BBox",
            "reference_type": "Reference",
            "ranking": "Ranking",
            "rank_changes_vs_human": "Changes vs human",
        }
    )[["Dataset", "BBox", "Reference", "Ranking", "Changes vs human"]]

    summary = [
        "Her pseudo referans GT-bbox koşulunda kendi üreticisine 1,000 IoU verir. Bu matematiksel özdeşlik başarı değil, deneyin pozitif kontrolüdür.",
        "YOLO bbox koşulunda da her pseudo referans kendi öğretmen modelini en yüksek skora taşır. Plane için öğretmen avantajı 0,113–0,138; Small Vehicle için 0,179–0,299 IoU aralığındadır.",
        "İnsan referansındaki model sıralaması pseudo referansa göre değişebilir. Bu nedenle model seçimi, referansı üreten model ailesine bağımlı hâle gelir.",
        "SAM3 Small Vehicle öğretmen çıktılarının 5.345/12.051'i (%44,4) boştur. SAM3'ün kendi pseudo referansında yüksek görünmesi bağımsız doğruluk kanıtı değildir.",
    ]
    definitions = [
        "Kapsam: iSAID Plane'de 5.447, iSAID Small Vehicle'da 12.051 instance; her veri setinde 512 görüntü ve dört eşit 128 görüntülük stratum.",
        "Sabitler: görüntüler, instance'lar, insan GT bbox istemleri, seed 42 YOLO bbox istemleri ve SAM1/2/3 tahminleri aynıdır. Değişen tek değerlendirme girdisi referans maskedir.",
        "Avg IoU her instance için ayrı hesaplanır ve instance'lar eşit ağırlıkla ortalanır. Büyük nesneler sonucu piksel alanıyla baskılamaz.",
        "Pseudo referans, SAM1/SAM2/SAM3'ün insan GT bbox istemiyle ürettiği maskedir. Bu nedenle insan lokalizasyon bilgisi korunur; deney maske sınırı ve model-öğretmen uyumuna odaklanır.",
        "Scene-clustered bootstrap güven aralıkları aynı kaynak sahneden gelen crop'ların bağımsız kabul edilmesini önler.",
    ]
    limitations = [
        "Bu çalışma pseudo etiketlerin eğitimde yararsız olduğunu kanıtlamaz. Test referansının bağımsız olmadığı durumda değerlendirme ve model sıralamasının bozulabildiğini gösterir.",
        "GT-bbox diagonal hücreleri tautological identity kontrolüdür; makalede ana performans sonucu olarak kullanılmamalıdır.",
        "YOLO-bbox sonuçları daha güçlü kanıttır, çünkü öğrenci istemi öğretmen referans isteminden farklıdır; ancak referans ve aday model aynı mimari aileden olduğu için hata korelasyonu hâlâ beklenir.",
        "SAM3 Small Vehicle boş maskeleri ayrı bir failure mode'dur. Boş referanslarda hem boş öğrenci tahmini 1,0 alabilir; boş oranı olmadan ortalama skor yanıltıcıdır.",
        "İki sınıf ve tek remote-sensing veri seti ailesiyle sınırlıyız. Sonuç genellenebilirlik iddiası değil, kontrollü bir ölçüm geçerliliği uyarısıdır.",
    ]
    recommendations = [
        "Ana model karşılaştırmasını insan referansı üzerinde raporla; pseudo referans sonuçlarını duyarlılık analizi olarak ayrı göster.",
        "Pseudo referans üreticisini, sürümünü, istemini, boş maske oranını ve post-processing adımlarını açıkça raporla.",
        "Referans üreticisiyle aynı model ailesini değerlendirirken diagonal/self-reference sonucunu performans tablosundan ayır veya açıkça identity control olarak işaretle.",
        "Mümkünse küçük, kör ve bağımsız bir insan audit alt kümesi kullan; model sıralamasının bu alt kümede korunup korunmadığını kontrol et.",
        "Bildiri ana figürü olarak model–referans matrisini, ana istatistik olarak pseudo−human paired IoU değişimini ve güven aralığını kullan.",
    ]

    markdown_path = output_dir / "sam_teacher_pseudo_reference_comparison.md"
    markdown_lines = [
        "# SAM1/SAM2/SAM3 Pseudo Referans Karşılaştırması",
        "",
        "## Teknik Özet",
        "",
        *[f"- {item}" for item in summary],
        "",
        "## Kapsam ve Tanımlar",
        "",
        *[f"- {item}" for item in definitions],
        "",
    ]
    for dataset_id in DATASETS:
        markdown_lines.extend([f"## {DATASET_LABELS[dataset_id]}", ""])
        for bbox in ("gt_bbox", "yolo_bbox"):
            markdown_lines.extend(
                [
                    f"### {'GT bbox' if bbox == 'gt_bbox' else 'YOLO bbox'} Overall Avg IoU",
                    "",
                    _markdown_table(matrices[(dataset_id, bbox)]),
                    "",
                ]
            )
    markdown_lines.extend(
        [
            "## Referanslar Arası Anlaşma",
            "",
            _markdown_table(agreements),
            "",
            "## Boş Maske Denetimi",
            "",
            _markdown_table(empties),
            "",
            "## Model Sıralamaları",
            "",
            _markdown_table(rankings_display),
            "",
            "## Sınırlılıklar",
            "",
            *[f"- {item}" for item in limitations],
            "",
            "## Önerilen Raporlama Protokolü",
            "",
            *[f"- {item}" for item in recommendations],
            "",
        ]
    )
    markdown_path.write_text("\n".join(markdown_lines), encoding="utf-8")

    docx_path = output_dir / "sam_teacher_pseudo_reference_comparison_colored.docx"
    document = Document()
    _configure_docx(document)
    document.add_heading("SAM1/SAM2/SAM3 Pseudo Referans Karşılaştırması", 0)
    document.add_heading("Teknik Özet", level=1)
    _docx_bullets(document, summary)
    document.add_heading("Kapsam ve Metrik Tanımları", level=1)
    _docx_bullets(document, definitions)
    document.add_heading("Model–Referans Matrisi", level=1)
    document.add_paragraph(
        "Satırlar değerlendirilen modeli, sütunlar referans kaynağını gösterir. "
        "Pseudo diagonalinin yükselmesi self-reference etkisini görünür kılar."
    )
    document.add_picture(str(matrix_figure), width=Inches(9.7))
    document.add_heading("İnsan Referansına Göre Skor Değişimi", level=1)
    document.add_paragraph(
        "Noktalar pseudo−human paired Avg IoU farkını, çizgiler kaynak sahne "
        "kümeli %95 bootstrap güven aralığını gösterir."
    )
    document.add_picture(str(effect_figure), width=Inches(9.7))
    for dataset_id in DATASETS:
        _docx_table(
            document,
            matrices[(dataset_id, "gt_bbox")],
            f"{DATASET_LABELS[dataset_id]} · GT bbox Overall Avg IoU",
        )
        _docx_table(
            document,
            matrices[(dataset_id, "yolo_bbox")],
            f"{DATASET_LABELS[dataset_id]} · YOLO bbox Overall Avg IoU",
        )
    _docx_table(document, metrics, "Her Koşuldaki En Yüksek Overall Metrikler")
    _docx_table(document, agreements, "Referanslar Arası Instance Anlaşması")
    _docx_table(document, empties, "Öğretmen Boş Maske Denetimi")
    _docx_table(document, rankings_display, "Referansa Göre Model Sıralaması")
    _docx_table(document, strata, "Stratum Bazında YOLO Teacher Advantage")
    document.add_heading("Sınırlılıklar ve Sağlamlık Kontrolleri", level=1)
    _docx_bullets(document, limitations)
    document.add_heading("Önerilen Raporlama Protokolü", level=1)
    _docx_bullets(document, recommendations)
    document.save(docx_path)

    pdf_path = output_dir / "sam_teacher_pseudo_reference_comparison_colored.pdf"
    with PdfPages(pdf_path) as pdf:
        _pdf_text_page(pdf, "SAM1/SAM2/SAM3 Pseudo Referans Karşılaştırması", summary)
        _pdf_text_page(pdf, "Kapsam ve Metrik Tanımları", definitions)
        _pdf_image_page(
            pdf,
            "Model–Referans IoU Matrisi",
            matrix_figure,
            "Satır: değerlendirilen model. Sütun: değerlendirme referansı. Diagonal hücreler self-reference kontrolüdür.",
        )
        _pdf_image_page(
            pdf,
            "İnsan Referansına Göre Skor Değişimi",
            effect_figure,
            "Pseudo−human paired Avg IoU farkı; kaynak sahne kümeli %95 bootstrap güven aralığı.",
        )
        for dataset_id in DATASETS:
            _pdf_table_page(
                pdf,
                DATASET_LABELS[dataset_id],
                [
                    ("GT bbox · Overall Avg IoU", matrices[(dataset_id, "gt_bbox")]),
                    ("YOLO bbox · Overall Avg IoU", matrices[(dataset_id, "yolo_bbox")]),
                ],
            )
        _pdf_table_page(
            pdf,
            "Referans Kalitesi ve Anlaşma",
            [("Boş maske denetimi", empties), ("Referanslar arası IoU", agreements)],
        )
        _pdf_table_page(pdf, "Model Sıralaması", [("Overall sıralamalar", rankings_display)])
        _pdf_text_page(pdf, "Sınırlılıklar ve Sağlamlık Kontrolleri", limitations)
        _pdf_text_page(pdf, "Önerilen Raporlama Protokolü", recommendations)

    return {"markdown": markdown_path, "docx": docx_path, "pdf": pdf_path}
